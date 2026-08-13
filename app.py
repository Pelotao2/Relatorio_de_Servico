import streamlit as st
import streamlit.components.v1 as components
import pandas as pd
import psycopg2
from psycopg2.extras import RealDictCursor
import os
from dotenv import load_dotenv
import time
import html
import json
import re
import difflib
import base64
import io
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Cm
from logica import (
    extrair_trecho, melhor_correspondencia_nome, extrair_viatura,
    interpretar_guarnicao_por_voz, calcular_km_rodado, status_troca_oleo,
    formatar_numero_fiscalizacao
)

# Mapa interativo (ponto + área) e conversão de coordenadas UTM — opcionais:
# se ainda não estiverem no requirements.txt, o app continua funcionando
# normalmente, só sem o mapa/UTM (mostra um aviso em vez de quebrar).
try:
    import folium
    from streamlit_folium import st_folium
    from folium.plugins import Draw
    MAPA_DISPONIVEL = True
except Exception:
    MAPA_DISPONIVEL = False

try:
    import utm as utm_lib
    UTM_DISPONIVEL = True
except Exception:
    UTM_DISPONIVEL = False

try:
    import pypdf
    PDF_LEITURA_DISPONIVEL = True
except Exception:
    PDF_LEITURA_DISPONIVEL = False

try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_DISPONIVEL = True
except Exception:
    OCR_DISPONIVEL = False

def exibir_logo(caminho, width):
    """Mostra uma imagem de logo sem derrubar o app se o arquivo não existir
    no ambiente (ex.: esquecido de subir pro GitHub na nuvem)."""
    try:
        st.image(caminho, width=width)
    except Exception:
        st.caption("🌿 (logo não encontrado neste ambiente)")

# --- SISTEMA DE AUTENTICAÇÃO ---
# Dicionário de usuários cadastrados (Altere/adicione os logins e senhas desejados aqui)
USUARIOS_PERMITIDOS = {
    "agente.pelotao": "SenhaPelotao2026",
    "comandante.123": "ComandoSeguro!#",
    "admin": "AdminNeon789"
}

# Inicializa as variáveis de controle de sessão se não existirem
if "autenticado" not in st.session_state:
    st.session_state["autenticado"] = False
if "usuario_conectado" not in st.session_state:
    st.session_state["usuario_conectado"] = ""
if "unidade_operacional" not in st.session_state:
    st.session_state["unidade_operacional"] = None

# Função que valida as credenciais
def realizar_login(usuario, senha):
    if usuario in USUARIOS_PERMITIDOS and USUARIOS_PERMITIDOS[usuario] == senha:
        st.session_state["autenticado"] = True
        st.session_state["usuario_conectado"] = usuario
        st.rerun()
    else:
        st.error("Usuário ou senha incorretos. Tente novamente.")

# --- PASSO 0: ESCOLHA DA UNIDADE (antes até da senha) ---
# Isola operacionalmente as duas unidades: cada uma só enxerga/mexe nos próprios
# serviços no Formulário de Serviço. O Painel Estratégico (Adm) continua vendo as duas.
if st.session_state["unidade_operacional"] is None:
    col_logo_login, col_vazio_login = st.columns([1, 3])
    with col_logo_login:
        exibir_logo("logo_pantanal_folhas.png", 180)
    st.title(" Sistema de Controle de Produtividade")
    st.subheader("Selecione sua unidade para continuar")
    col_un_a, col_un_b = st.columns(2)
    with col_un_a:
        if st.button("2º Pel Miranda", use_container_width=True, type="primary"):
            st.session_state["unidade_operacional"] = "2º Pel Miranda"
            st.rerun()
    with col_un_b:
        if st.button("GPM Barra", use_container_width=True, type="primary"):
            st.session_state["unidade_operacional"] = "GPM Barra"
            st.rerun()
    st.caption("O Painel Estratégico (Adm) continua com visão das duas unidades, após o login.")
    st.stop()

# Se o usuário NÃO estiver autenticado, exibe a tela de login e para a execução do app aqui
if not st.session_state["autenticado"]:
    st.title("🔒 Sistema de Controle de Produtividade")
    st.caption(f"Unidade selecionada: **{st.session_state['unidade_operacional']}**")
    if st.button("↩️ Trocar unidade"):
        st.session_state["unidade_operacional"] = None
        st.rerun()
    st.subheader("Efetue o login para acessar o formulário")
    
    with st.form("formulario_login"):
        user_input = st.text_input("Usuário / ID Funcional", placeholder="Ex: agente.nome")
        pass_input = st.text_input("Senha", type="password", placeholder="Digite sua senha")
        botao_entrar = st.form_submit_button("Acessar Sistema")
        
        if botao_entrar:
            realizar_login(user_input, pass_input)
            
    # Interrompe o carregamento do restante do código (Bloqueia o acesso)
    st.stop()

# --- FIM DO SISTEMA DE AUTENTICAÇÃO ---
if "patrulhamento_terrestre_list" not in st.session_state:
    st.session_state["patrulhamento_terrestre_list"] = []
if "patrulhamento_fluvial_list" not in st.session_state:
    st.session_state["patrulhamento_fluvial_list"] = []
if "capturas_animais_list" not in st.session_state:
    st.session_state["capturas_animais_list"] = []
if "apreensoes_list" not in st.session_state:
    st.session_state["apreensoes_list"] = []
if "prolepse_list" not in st.session_state:
    st.session_state["prolepse_list"] = []
if "relatorio_id_atual" not in st.session_state:
    st.session_state["relatorio_id_atual"] = None
if "guarnicao_carregada_key" not in st.session_state:
    st.session_state["guarnicao_carregada_key"] = None
if "editando_idx_animal" not in st.session_state:
    st.session_state["editando_idx_animal"] = None
if "editando_idx_terrestre" not in st.session_state:
    st.session_state["editando_idx_terrestre"] = None
if "editando_idx_fluvial" not in st.session_state:
    st.session_state["editando_idx_fluvial"] = None
if "editando_idx_apreensao" not in st.session_state:
    st.session_state["editando_idx_apreensao"] = None
if "editando_idx_prolepse" not in st.session_state:
    st.session_state["editando_idx_prolepse"] = None
if "cautela_itens_temp" not in st.session_state:
    st.session_state["cautela_itens_temp"] = []
if "doc_geracao" not in st.session_state:
    st.session_state["doc_geracao"] = 0

# Configuração da página (Modo Largo para Computadores do Quartel)
st.set_page_config(page_title="Sistema de Produtividade - Pelotão", page_icon="icone_pelotao_512.png", layout="wide")

# CSS de impressão: some com cabeçalho, barra de abas e botões apenas na hora
# de imprimir/salvar em PDF, deixando o relatório mais limpo no papel.
st.markdown("""
    <style>
    .print-only-text { display: none; }
    .ajuda-tela { font-size: 0.85em; color: #888; }
    @media print {
        header[data-testid="stHeader"],
        div[data-testid="stToolbar"],
        section[data-testid="stSidebar"],
        div[data-baseweb="tab-list"],
        [data-testid="stTabs"] [role="tablist"],
        .stTabs > div:first-child,
        [role="tablist"],
        .stButton, .stDownloadButton, #MainMenu, footer {
            display: none !important;
        }
        [data-testid="stTextArea"], textarea {
            display: none !important;
        }
        [data-testid="stDialog"],
        div[role="dialog"],
        [data-testid="stModal"],
        div[data-baseweb="modal"] {
            display: none !important;
        }
        [data-testid="stAlert"] {
            display: none !important;
        }
        [data-testid="stIconMaterial"],
        [data-testid="stIcon"] {
            display: none !important;
        }
        .ajuda-tela {
            display: none !important;
        }
        .print-only-text {
            display: block !important;
            white-space: pre-wrap;
            margin-bottom: 12px;
        }
        body, p, span, div, label, .print-only-text, textarea, input {
            font-family: Arial, Helvetica, sans-serif !important;
            font-size: 12pt !important;
        }
        h1 { font-size: 18pt !important; }
        h2, h3 { font-size: 15pt !important; }
        h4, h5, h6, .stMarkdown strong { font-size: 13pt !important; }
    }
    </style>
""", unsafe_allow_html=True)
def ajuda(texto):
    """Texto de orientação/ajuda de tela — nunca aparece na impressão
    (diferente de st.caption, que aparece — use para conteúdo real do relatório)."""
    st.markdown(f'<div class="ajuda-tela">{texto}</div>', unsafe_allow_html=True)
# ==========================================
# CONEXÃO COM O BANCO DE DADOS (NEON PG)
# ==========================================
load_dotenv()
DATABASE_URL = os.environ.get("DATABASE_URL")

def init_connection():
    if not DATABASE_URL:
        return None
    try:
        return psycopg2.connect(DATABASE_URL, cursor_factory=RealDictCursor)
    except Exception as e:
        return None

def buscar_troca_oleo(placa):
    """Busca o KM da última troca de óleo registrada para a viatura (placa).
    Retorna 0 se a viatura ainda não tiver nenhum registro."""
    conn = init_connection()
    if not conn:
        return 0
    try:
        cur = conn.cursor()
        cur.execute("SELECT ultima_troca_km FROM viaturas WHERE placa = %s;", (placa,))
        registro = cur.fetchone()
        cur.close()
        conn.close()
        return registro["ultima_troca_km"] if registro else 0
    except Exception:
        return 0

def salvar_troca_oleo(placa, km):
    """Grava (insere ou atualiza) o KM da última troca de óleo de uma viatura."""
    conn = init_connection()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO viaturas (placa, ultima_troca_km) VALUES (%s, %s) "
            "ON CONFLICT (placa) DO UPDATE SET ultima_troca_km = EXCLUDED.ultima_troca_km;",
            (placa, km)
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception:
        return False

def criar_cautela(unidade, destinatario, data_cautela, prazo, prazo_indefinido, itens):
    """Cria uma nova cautela de armamento/munição com um ou mais itens."""
    conn = init_connection()
    if not conn:
        return None, "Sem conexão com o banco de dados."
    try:
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO cautelas (unidade, destinatario, data_cautela, prazo, prazo_indefinido, itens, status) "
            "VALUES (%s, %s, %s, %s, %s, %s, %s) RETURNING id;",
            (unidade, destinatario, data_cautela, prazo, prazo_indefinido, json.dumps(itens, ensure_ascii=False), "Em Aberto")
        )
        novo_id = cur.fetchone()["id"]
        conn.commit()
        cur.close()
        conn.close()
        return novo_id, None
    except Exception as e:
        conn.close()
        return None, str(e)

def listar_cautelas_abertas(unidade):
    """Lista todas as cautelas 'Em Aberto' da unidade — aparecem em todos os
    relatórios até que o material seja entregue de volta."""
    conn = init_connection()
    if not conn:
        return []
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM cautelas WHERE unidade = %s AND status = %s ORDER BY id DESC;",
            (unidade, "Em Aberto")
        )
        registros = cur.fetchall()
        cur.close()
        conn.close()
        return registros
    except Exception:
        return []

def entregar_cautela(cautela_id, observacao):
    """Marca uma cautela como entregue (material devolvido), encerrando-a."""
    conn = init_connection()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute(
            "UPDATE cautelas SET status = %s, observacao_entrega = %s, data_entrega = CURRENT_DATE WHERE id = %s;",
            ("Entregue", observacao, cautela_id)
        )
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception:
        return False

# --- Designação oficial da unidade, usada no número do Relatório de Fiscalização ---
# Nº <sequencial>/<designação>/<ano>. Ajuste se a GPM Barra usar uma designação diferente.
DESIGNACAO_UNIDADE_FISCALIZACAO = {
    "2º Pel Miranda": "2ºPEL/2ªCIA/1ºBPMA/CPAMB",
    "GPM Barra": "2ºPEL/2ªCIA/1ºBPMA/CPAMB"
}

def proximo_numero_fiscalizacao(unidade, ano):
    """Calcula o próximo número sequencial do Relatório de Fiscalização, por
    unidade e por ano (reinicia a cada ano, como no modelo oficial)."""
    conn = init_connection()
    if not conn:
        return 1
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT COALESCE(MAX(sequencial), 0) + 1 AS proximo FROM relatorios_fiscalizacao WHERE unidade = %s AND ano = %s;",
            (unidade, ano)
        )
        return cur.fetchone()["proximo"]
    except Exception:
        return 1
    finally:
        conn.close()

def dms_para_decimal(graus, minutos, segundos, hemisferio):
    """Converte Graus/Minutos/Segundos para Graus Decimais (float)."""
    decimal = graus + minutos / 60 + segundos / 3600
    if hemisferio in ("S", "W"):
        decimal = -decimal
    return decimal

def decimal_para_dms(decimal, eh_latitude=True):
    """Converte Graus Decimais para uma tupla (graus, minutos, segundos, hemisfério)."""
    hemisferio = ("N" if decimal >= 0 else "S") if eh_latitude else ("E" if decimal >= 0 else "W")
    decimal_abs = abs(decimal)
    graus = int(decimal_abs)
    resto = (decimal_abs - graus) * 60
    minutos = int(resto)
    segundos = (resto - minutos) * 60
    return graus, minutos, segundos, hemisferio

def campo_geolocalizacao(key_prefix, permitir_area=True):
    """Campo de localização geográfica com 3 formatos de entrada (Graus Decimais,
    GMS e UTM) e mapa interativo (marcar ponto clicando e, opcionalmente, desenhar
    uma área/polígono). Retorna (latitude, longitude, coordenadas_texto, area_geojson).
    Todos podem vir None/"" se nada foi preenchido ainda."""
    st.markdown("###### 📍 Localização Geográfica")
    formato = st.radio(
        "Formato de entrada das coordenadas",
        ["Graus Decimais (DD)", "Graus, Minutos e Segundos (GMS)", "UTM"],
        key=f"{key_prefix}_formato", horizontal=True
    )

    lat, lon = None, None

    if formato == "Graus Decimais (DD)":
        col_dd1, col_dd2 = st.columns(2)
        with col_dd1:
            lat = st.number_input("Latitude (DD)", value=0.0, format="%.6f", step=0.000001, key=f"{key_prefix}_lat_dd")
        with col_dd2:
            lon = st.number_input("Longitude (DD)", value=0.0, format="%.6f", step=0.000001, key=f"{key_prefix}_lon_dd")

    elif formato == "Graus, Minutos e Segundos (GMS)":
        st.caption("Latitude")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            lat_g = st.number_input("Graus", min_value=0, max_value=90, step=1, key=f"{key_prefix}_lat_g")
        with c2:
            lat_m = st.number_input("Minutos", min_value=0, max_value=59, step=1, key=f"{key_prefix}_lat_m")
        with c3:
            lat_s = st.number_input("Segundos", min_value=0.0, max_value=59.999, step=0.1, key=f"{key_prefix}_lat_s")
        with c4:
            lat_h = st.selectbox("Hemisfério", ["S", "N"], key=f"{key_prefix}_lat_h")
        st.caption("Longitude")
        c5, c6, c7, c8 = st.columns(4)
        with c5:
            lon_g = st.number_input("Graus", min_value=0, max_value=180, step=1, key=f"{key_prefix}_lon_g")
        with c6:
            lon_m = st.number_input("Minutos", min_value=0, max_value=59, step=1, key=f"{key_prefix}_lon_m")
        with c7:
            lon_s = st.number_input("Segundos", min_value=0.0, max_value=59.999, step=0.1, key=f"{key_prefix}_lon_s")
        with c8:
            lon_h = st.selectbox("Hemisfério", ["W", "E"], key=f"{key_prefix}_lon_h")
        lat = dms_para_decimal(lat_g, lat_m, lat_s, lat_h)
        lon = dms_para_decimal(lon_g, lon_m, lon_s, lon_h)

    else:  # UTM
        if not UTM_DISPONIVEL:
            st.warning("⚠️ Conversão UTM indisponível neste ambiente — adicione `utm` ao requirements.txt e reimplante o app.")
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            zona_utm = st.number_input("Zona UTM", min_value=1, max_value=60, value=21, step=1, key=f"{key_prefix}_utm_zona")
        with c2:
            hemisferio_utm = st.selectbox("Hemisfério", ["S", "N"], key=f"{key_prefix}_utm_hemisferio")
        with c3:
            easting = st.number_input("Easting (E)", value=0.0, format="%.2f", key=f"{key_prefix}_utm_e")
        with c4:
            northing = st.number_input("Northing (N)", value=0.0, format="%.2f", key=f"{key_prefix}_utm_n")
        if UTM_DISPONIVEL and easting and northing:
            try:
                lat, lon = utm_lib.to_latlon(easting, northing, int(zona_utm), northern=(hemisferio_utm == "N"))
            except Exception as e:
                st.error(f"Coordenada UTM inválida: {e}")

    coordenadas_texto = f"{lat:.6f}, {lon:.6f}" if lat is not None and lon is not None else ""
    area_geojson = None

    if not MAPA_DISPONIVEL:
        st.info("🗺️ Mapa indisponível neste ambiente — adicione `folium` e `streamlit-folium` ao requirements.txt e reimplante o app para habilitar o mapa interativo.")
        return lat, lon, coordenadas_texto, area_geojson

    centro = [lat, lon] if (lat and lon) else [-20.24, -56.38]  # aprox. região de Miranda/MS
    mapa = folium.Map(location=centro, zoom_start=14 if (lat and lon) else 9)
    if lat and lon:
        icone_pino = folium.DivIcon(html='<div style="font-size:30px; line-height:30px; transform:translate(-50%,-100%);">📍</div>')
        folium.Marker([lat, lon], tooltip="Local marcado", icon=icone_pino).add_to(mapa)
    if permitir_area:
        Draw(
            export=False,
            draw_options={"marker": True, "polygon": True, "rectangle": True, "circle": False, "circlemarker": False, "polyline": False},
            edit_options={"edit": True},
        ).add_to(mapa)
    resultado_mapa = st_folium(mapa, height=350, key=f"{key_prefix}_mapa", use_container_width=True)

    if resultado_mapa:
        clique = resultado_mapa.get("last_clicked")
        if clique:
            lat, lon = clique["lat"], clique["lng"]
            coordenadas_texto = f"{lat:.6f}, {lon:.6f}"
            st.caption(f"📍 Ponto atualizado pelo clique no mapa: {coordenadas_texto} (o campo de coordenadas acima não se atualiza sozinho — use este valor ao registrar)")
        desenho = resultado_mapa.get("last_active_drawing")
        if desenho:
            area_geojson = json.dumps(desenho, ensure_ascii=False)
            st.caption("🔷 Área/ponto desenhado no mapa capturado com sucesso.")

    return lat, lon, coordenadas_texto, area_geojson

MESES_PT_PARA_NUM = {
    "janeiro": 1, "fevereiro": 2, "março": 3, "abril": 4, "maio": 5, "junho": 6,
    "julho": 7, "agosto": 8, "setembro": 9, "outubro": 10, "novembro": 11, "dezembro": 12
}

def parse_data_pt(texto_data):
    """Converte algo como '17 de julho de 2026' para um objeto date. Retorna
    None se não conseguir reconhecer o padrão (o usuário preenche na mão)."""
    m = re.search(r"(\d{1,2})\s*de\s*([a-zçãé]+)\s*de\s*(\d{4})", texto_data, re.IGNORECASE)
    if not m:
        return None
    dia, mes_nome, ano = m.groups()
    mes_num = MESES_PT_PARA_NUM.get(mes_nome.lower().strip())
    if not mes_num:
        return None
    try:
        return datetime(int(ano), mes_num, int(dia)).date()
    except Exception:
        return None

def extrair_texto_pdf(arquivo_pdf):
    """Extrai o texto de um PDF. Primeiro tenta o caminho rápido (texto nativo,
    selecionável). Se vier vazio/quase vazio — comum em documentos assinados
    digitalmente pelo gov.br, que costumam 'achatar' cada página em uma imagem —
    cai automaticamente para OCR (reconhecimento de texto na imagem da página)."""
    leitor = pypdf.PdfReader(arquivo_pdf)
    texto_completo = ""
    for pagina in leitor.pages:
        texto_completo += (pagina.extract_text() or "") + "\n"

    if len(texto_completo.strip()) >= 50:
        return texto_completo, "texto"

    if not OCR_DISPONIVEL:
        return texto_completo, "texto"

    # Fallback: OCR. Precisa reler o arquivo do início (pypdf já consumiu o ponteiro).
    arquivo_pdf.seek(0)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(arquivo_pdf.read())
        caminho_tmp = tmp.name
    try:
        paginas_imagem = convert_from_path(caminho_tmp, dpi=200)
        texto_ocr = ""
        for imagem_pagina in paginas_imagem:
            texto_ocr += pytesseract.image_to_string(imagem_pagina, lang="por") + "\n"
        return texto_ocr, "ocr"
    finally:
        try:
            os.remove(caminho_tmp)
        except Exception:
            pass

def extrair_dados_ordem_operacao(texto):
    """Tenta sugerir número, vigência, finalidade resumida e instruções de
    CADG/Kobotoolbox a partir do texto da Ordem de Operação. Qualquer campo não
    encontrado volta vazio/None — a pessoa preenche/confirma na tela, nunca é
    salvo sem revisão humana."""
    sugestao = {"numero": "", "data_inicio": None, "data_fim": None, "finalidade": "", "cadg_kobo": ""}

    m_num = re.search(r"ORDEM DE OPERA[ÇC][ÃA]O\s*N[ºo°]?\s*([^\n]+)", texto, re.IGNORECASE)
    if m_num:
        sugestao["numero"] = m_num.group(1).strip()

    m_ini = re.search(r"Data de In[ií]cio\s*:?\s*([^\n]+)", texto, re.IGNORECASE)
    if m_ini:
        sugestao["data_inicio"] = parse_data_pt(m_ini.group(1))

    m_fim = re.search(r"Data\s*(?:de\s*)?T[ée]rmino\s*:?\s*([^\n]+)", texto, re.IGNORECASE)
    if m_fim:
        sugestao["data_fim"] = parse_data_pt(m_fim.group(1))

    m_final = re.search(r"1\.\s*FINALIDADE\s*(.*?)(?:2\.\s*OBJETIVO|$)", texto, re.IGNORECASE | re.DOTALL)
    if m_final:
        sugestao["finalidade"] = re.sub(r"\s+", " ", m_final.group(1)).strip()[:600]

    m_presc = re.search(r"11\.\s*PRESCRI[ÇC][ÕO]ES DIVERSAS\s*(.*)", texto, re.IGNORECASE | re.DOTALL)
    if m_presc:
        sugestao["cadg_kobo"] = re.sub(r"\s+", " ", m_presc.group(1)).strip()[:2000]

    return sugestao

def montar_texto_ordem_servico(ordem):
    """Monta o texto final da Ordem de Serviço a partir dos dados confirmados
    da Ordem de Operação."""
    finalidade = ordem.get("finalidade") or "[finalidade não especificada]"
    texto = f"Em atendimento à Ordem de Operação {ordem.get('numero','')}, a equipe deverá realizar {finalidade}. "
    texto += "Ao final, a equipe deverá registrar o fato no CADG citando a Ordem em vigor"
    if ordem.get("cadg_kobo"):
        texto += f", observando ainda: {ordem['cadg_kobo']}"
    else:
        texto += " e também deverá lançar os dados no KoboToolbox."
    return texto

def salvar_ordem_operacao(dados, usuario):
    conn = init_connection()
    if not conn:
        return None, "Sem conexão com o banco."
    try:
        cur = conn.cursor()
        _pdf_bytes = psycopg2.Binary(dados["arquivo_pdf"]) if dados.get("arquivo_pdf") else None
        cur.execute(
            """INSERT INTO ordens_operacao
               (numero, data_inicio, data_fim, finalidade, cadg_kobo, status, criado_por, data_criacao, arquivo_pdf, arquivo_pdf_nome)
               VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s) RETURNING id;""",
            (dados["numero"], dados["data_inicio"], dados["data_fim"], dados["finalidade"],
             dados["cadg_kobo"], "Vigente", usuario, datetime.now(), _pdf_bytes, dados.get("arquivo_pdf_nome"))
        )
        novo_id = cur.fetchone()["id"]
        conn.commit()
        cur.close()
        conn.close()
        registrar_auditoria(novo_id, "ordem_operacao", "Criação", usuario)
        return novo_id, None
    except Exception as e:
        conn.close()
        return None, str(e)

def listar_ordens_operacao():
    conn = init_connection()
    if not conn:
        return []
    try:
        cur = conn.cursor()
        cur.execute("SELECT * FROM ordens_operacao ORDER BY data_inicio DESC;")
        registros = cur.fetchall()
        cur.close()
        conn.close()
        return registros
    except Exception:
        return []

def encerrar_ordem_manualmente(id_ordem, usuario):
    conn = init_connection()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute(
            "UPDATE ordens_operacao SET status = %s, encerrada_manualmente_em = %s, encerrada_por = %s WHERE id = %s;",
            ("Encerrada Manualmente", datetime.now(), usuario, id_ordem)
        )
        conn.commit()
        cur.close()
        conn.close()
        registrar_auditoria(id_ordem, "ordem_operacao", "Encerramento Manual", usuario)
        return True
    except Exception:
        return False

def registrar_auditoria(relatorio_id, tipo_relatorio, acao, usuario):
    """Grava uma linha no log de auditoria (quem fez o quê e quando).
    Nunca interrompe o fluxo principal se falhar — ex.: se a tabela
    'log_auditoria' ainda não existir no banco, apenas ignora em silêncio."""
    if not relatorio_id:
        return
    conn = init_connection()
    if not conn:
        return
    try:
        cur = conn.cursor()
        cur.execute(
            "INSERT INTO log_auditoria (relatorio_id, tipo_relatorio, acao, usuario, data_hora) VALUES (%s, %s, %s, %s, %s);",
            (relatorio_id, tipo_relatorio, acao, usuario or "desconhecido", datetime.now())
        )
        conn.commit()
        cur.close()
        conn.close()
    except Exception:
        try:
            conn.close()
        except Exception:
            pass

def salvar_relatorio_fiscalizacao(dados: dict, id_existente=None, usuario_atual=""):
    """Cria ou atualiza um Relatório de Fiscalização Ambiental. Em edições,
    registra quem alterou e quando (auditoria)."""
    conn = init_connection()
    if not conn:
        return None, "Sem conexão com o banco de dados."
    try:
        cur = conn.cursor()
        if id_existente:
            dados = dict(dados)
            dados["editado_por"] = usuario_atual
            dados["data_edicao"] = datetime.now()
            campos_imutaveis = {"numero", "sequencial", "ano", "unidade", "criado_por"}
            dados = {k: v for k, v in dados.items() if k not in campos_imutaveis}
            sets = ", ".join(f"{col} = %s" for col in dados.keys())
            valores = list(dados.values()) + [id_existente]
            cur.execute(f"UPDATE relatorios_fiscalizacao SET {sets} WHERE id = %s;", valores)
            novo_id = id_existente
        else:
            dados = dict(dados)
            dados["criado_por"] = usuario_atual
            colunas = ", ".join(dados.keys())
            placeholders = ", ".join(["%s"] * len(dados))
            cur.execute(f"INSERT INTO relatorios_fiscalizacao ({colunas}) VALUES ({placeholders}) RETURNING id;", list(dados.values()))
            novo_id = cur.fetchone()["id"]
        conn.commit()
        cur.close()
        conn.close()
        registrar_auditoria(novo_id, "fiscalizacao", "Edição" if id_existente else "Criação", usuario_atual)
        return novo_id, None
    except Exception as e:
        conn.close()
        return None, str(e)

def buscar_relatorios_fiscalizacao(termo, unidade=None):
    """Busca Relatórios de Fiscalização por número ou nome do autuado.
    Se `unidade` for None, busca nas duas unidades (uso do Painel Admin)."""
    conn = init_connection()
    if not conn:
        return []
    try:
        cur = conn.cursor()
        termo_like = f"%{termo}%"
        if unidade:
            cur.execute(
                "SELECT * FROM relatorios_fiscalizacao WHERE unidade = %s AND (numero ILIKE %s OR nome_autuado ILIKE %s) ORDER BY id DESC LIMIT 30;",
                (unidade, termo_like, termo_like)
            )
        else:
            cur.execute(
                "SELECT * FROM relatorios_fiscalizacao WHERE numero ILIKE %s OR nome_autuado ILIKE %s ORDER BY id DESC LIMIT 30;",
                (termo_like, termo_like)
            )
        registros = cur.fetchall()
        cur.close()
        conn.close()
        return registros
    except Exception:
        return []

def excluir_relatorio_fiscalizacao(id_relatorio):
    """Exclui definitivamente um Relatório de Fiscalização — uso restrito ao Admin."""
    conn = init_connection()
    if not conn:
        return False
    try:
        cur = conn.cursor()
        cur.execute("DELETE FROM relatorios_fiscalizacao WHERE id = %s;", (id_relatorio,))
        conn.commit()
        cur.close()
        conn.close()
        return True
    except Exception:
        return False

def buscar_relatorio_em_andamento(unidade, comandante):
    """Procura, para a guarnição informada, um relatório com status 'Em Andamento'.
    Usado para retomar automaticamente o serviço do dia anterior."""
    conn = init_connection()
    if not conn:
        return None
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM relatorios_servico WHERE unidade = %s AND comandante = %s AND status = %s ORDER BY id DESC LIMIT 1;",
            (unidade, comandante, "Em Andamento")
        )
        registro = cur.fetchone()
        cur.close()
        conn.close()
        return registro
    except Exception:
        return None

def _docx_shading(cell, cor_hex):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), cor_hex)
    cell._tc.get_or_add_tcPr().append(shd)

def _docx_secao_titulo(doc, texto):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    tabela = doc.add_table(rows=1, cols=1)
    tabela.style = "Table Grid"
    celula = tabela.rows[0].cells[0]
    _docx_shading(celula, "00D25F")
    p = celula.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(14)

def _docx_campo(doc, campos):
    """campos: lista de tuplas (numero_ou_None, rotulo, valor), lado a lado na mesma linha."""
    tabela = doc.add_table(rows=2, cols=len(campos))
    tabela.style = "Table Grid"
    for idx, (numero, rotulo, valor) in enumerate(campos):
        texto_rotulo = f"{numero}   {rotulo}" if numero else rotulo
        celula_label = tabela.rows[0].cells[idx]
        _docx_shading(celula_label, "D7E3BC")
        p_label = celula_label.paragraphs[0]
        run_label = p_label.add_run(texto_rotulo)
        run_label.bold = True
        run_label.font.size = Pt(9)
        p_val = tabela.rows[1].cells[idx].paragraphs[0]
        run_val = p_val.add_run(str(valor or ""))
        run_val.font.size = Pt(10)

def _docx_add_page_number_field(paragraph):
    """Insere o campo dinâmico de número de página (equivalente a PAGE do Word)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "PAGE"
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

def _docx_add_numpages_field(paragraph):
    """Insere o campo dinâmico de total de páginas (equivalente a NUMPAGES do Word)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    run = paragraph.add_run()
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "NUMPAGES"
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)

# Texto do cabeçalho oficial por unidade — confirme/ajuste a designação da GPM Barra.
DESIGNACAO_CABECALHO_FISCALIZACAO = {
    "2º Pel Miranda": "2º PELOTÃO DE PM AMBIENTAL/2°CPMA/1°BPMA/CPAMB/MIRANDA",
    "GPM Barra": "2º PELOTÃO DE PM AMBIENTAL/2°CPMA/1°BPMA/CPAMB/BARRA"
}
RODAPE_INSTITUCIONAL_FISCALIZACAO = (
    "Rua Lima Félix nº 175 – Jd. Veraneio – Pq. das Nações Indígenas - Cep 79037-109, "
    "Campo Grande–MS Fones: (67) 3357-1500 - e-mail: pmams_p1@hotmail.com"
)

def gerar_docx_fiscalizacao(dados):
    """Gera o documento .docx do Relatório de Fiscalização Ambiental, fiel ao
    modelo oficial (cabeçalho com logo e designação, formulários numerados,
    fatos, fotos, multa, providências, assinatura e rodapé institucional)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    doc = Document()
    for secao in doc.sections:
        secao.left_margin = Cm(2)
        secao.right_margin = Cm(2)

        header = secao.header
        p_logo = header.paragraphs[0]
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p_logo.add_run().add_picture("logo_pma_oficial.png", width=Cm(9))
        except Exception:
            pass
        p_designacao = header.add_paragraph()
        p_designacao.alignment = WD_ALIGN_PARAGRAPH.CENTER
        designacao_cab = DESIGNACAO_CABECALHO_FISCALIZACAO.get(dados.get("unidade", ""), "")
        r_desig = p_designacao.add_run(designacao_cab)
        r_desig.bold = True
        r_desig.font.size = Pt(11)

        footer = secao.footer
        p_end = footer.paragraphs[0]
        p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_end = p_end.add_run(RODAPE_INSTITUCIONAL_FISCALIZACAO)
        r_end.font.size = Pt(8)
        p_pag = footer.add_paragraph()
        p_pag.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_pag1 = p_pag.add_run("Página "); r_pag1.font.size = Pt(8)
        _docx_add_page_number_field(p_pag)
        r_pag2 = p_pag.add_run(" de "); r_pag2.font.size = Pt(8)
        _docx_add_numpages_field(p_pag)

    titulo = doc.add_paragraph()
    titulo.alignment = 1
    r = titulo.add_run("RELATÓRIO DE FISCALIZAÇÃO AMBIENTAL")
    r.bold = True
    r.font.size = Pt(13)

    numero_par = doc.add_paragraph()
    numero_par.alignment = 1
    r2 = numero_par.add_run(f"Nº {dados.get('numero','')}")
    r2.bold = True

    p_int = doc.add_paragraph()
    p_int.add_run("Interessado: ").bold = True
    p_int.add_run(dados.get("interessado", ""))

    _docx_secao_titulo(doc, "DO AUTUADO/FISCALIZADO")
    _docx_campo(doc, [("01", "Nome/Nome Empresarial", dados.get("nome_autuado", ""))])
    _docx_campo(doc, [("02", "CPF/CNPJ", dados.get("cpf_cnpj", "")), ("03", "RG/Insc. Estadual", dados.get("rg_ie", ""))])
    _docx_campo(doc, [("04", "Endereço Completo", dados.get("endereco", ""))])

    _docx_secao_titulo(doc, "DA INFRAÇÃO/FISCALIZAÇÃO")
    _docx_campo(doc, [("05", "Local", dados.get("local_fiscalizacao", "")), ("06", "Data", dados.get("data_fiscalizacao", ""))])
    _docx_campo(doc, [
        ("07", "Coord. Geográfica", dados.get("coordenadas", "")),
        ("08", "Município", dados.get("municipio", "")),
        ("09", "Telefone", dados.get("telefone", ""))
    ])

    _docx_secao_titulo(doc, "LEGISLAÇÃO APLICÁVEL")
    _docx_campo(doc, [("10", "Legislação", dados.get("legislacao", ""))])

    _docx_secao_titulo(doc, "FORMULÁRIOS IMASUL")
    _docx_campo(doc, [("", "Auto de Infração nº:", dados.get("auto_infracao_nr", "")), ("", "Laudo de Constatação nº:", dados.get("laudo_constatacao_nr", ""))])
    _docx_campo(doc, [("", "Termo de Paralisação nº:", dados.get("termo_paralisacao_nr", "")), ("", "Notificação nº:", dados.get("notificacao_nr", ""))])
    _docx_campo(doc, [("", "Folhas Complementares (quantidade):", dados.get("folhas_complementares", "")), ("", "BO CADG nº:", dados.get("bo_cadg_nr", ""))])

    doc.add_paragraph()
    p1 = doc.add_paragraph()
    p1.add_run("1. DOS FATOS").bold = True
    p2 = doc.add_paragraph()
    p2.add_run("1. HISTÓRICO").bold = True
    doc.add_paragraph(dados.get("fatos_historico", ""))

    try:
        fotos = json.loads(dados.get("fotos") or "[]")
    except Exception:
        fotos = []
    for foto in fotos:
        try:
            img_bytes = base64.b64decode(foto["dados_base64"])
            doc.add_picture(io.BytesIO(img_bytes), width=Cm(9))
            legenda = doc.add_paragraph(foto.get("legenda", ""))
            legenda.alignment = 1
        except Exception:
            pass

    p3 = doc.add_paragraph()
    p3.add_run("2. DO VALOR DA MULTA:").bold = True
    doc.add_paragraph(dados.get("valor_multa_texto", ""))

    p4 = doc.add_paragraph()
    p4.add_run("3. DAS PROVIDÊNCIAS ADMINISTRATIVAS").bold = True
    doc.add_paragraph(dados.get("providencias", ""))

    doc.add_paragraph()
    p_local = doc.add_paragraph()
    p_local.alignment = 1
    p_local.add_run(f"{dados.get('municipio_assinatura','Miranda (MS)')}, {dados.get('data_assinatura','')}.")

    p_ass1 = doc.add_paragraph()
    p_ass1.alignment = 1
    p_ass1.add_run(dados.get("relator", "")).bold = True
    p_ass2 = doc.add_paragraph()
    p_ass2.alignment = 1
    p_ass2.add_run(dados.get("cargo_relator", "Cmt. da GU Ambiental/Relator"))
    p_ass3 = doc.add_paragraph()
    p_ass3.alignment = 1
    p_ass3.add_run(f"Mat. {dados.get('matricula_relator','')}")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def listar_relatorios_em_andamento(unidade):
    """Lista todos os relatórios 'Em Andamento' de uma unidade — usado no painel
    'GU Serviço' no topo do formulário, já isolado por unidade operacional."""
    conn = init_connection()
    if not conn:
        return []
    try:
        cur = conn.cursor()
        cur.execute(
            "SELECT * FROM relatorios_servico WHERE unidade = %s AND status = %s ORDER BY id DESC;",
            (unidade, "Em Andamento")
        )
        registros = cur.fetchall()
        cur.close()
        conn.close()
        return registros
    except Exception:
        return []

def buscar_relatorios(unidade, termo):
    """Busca relatórios (qualquer status) por Nº do relatório ou nome do
    comandante, restrito à unidade operacional atual. Usado para localizar
    relatórios já concluídos e conferir/editar."""
    conn = init_connection()
    if not conn:
        return []
    try:
        cur = conn.cursor()
        termo_like = f"%{termo}%"
        cur.execute(
            "SELECT * FROM relatorios_servico WHERE unidade = %s AND (CAST(id AS TEXT) ILIKE %s OR comandante ILIKE %s) ORDER BY id DESC LIMIT 20;",
            (unidade, termo_like, termo_like)
        )
        registros = cur.fetchall()
        cur.close()
        conn.close()
        return registros
    except Exception:
        return []

def salvar_relatorio(colunas_valores: dict, id_existente=None):
    """Grava o relatório no banco: insere um novo registro (dia 1) ou atualiza
    um registro 'Em Andamento' já existente (dias 2 a 5). Retorna (id, erro)."""
    conn = init_connection()
    if not conn:
        return None, "Sem conexão com o banco de dados."
    try:
        cur = conn.cursor()
        if id_existente:
            # Campos "fundadores" do serviço: só são gravados na criação (Dia 1).
            # Nenhum salvamento posterior (autosave, edição de itens, progresso
            # do dia) pode alterá-los — evita que a contagem de dias/identidade
            # da guarnição seja corrompida por uma retomada ou salvamento parcial.
            campos_imutaveis_apos_criacao = {"data_inicial", "km_inicial", "unidade", "comandante"}
            colunas_valores = {k: v for k, v in colunas_valores.items() if k not in campos_imutaveis_apos_criacao}
            sets = ", ".join(f"{col} = %s" for col in colunas_valores.keys())
            valores = list(colunas_valores.values()) + [id_existente]
            cur.execute(f"UPDATE relatorios_servico SET {sets} WHERE id = %s;", valores)
            novo_id = id_existente
        else:
            colunas = ", ".join(colunas_valores.keys())
            placeholders = ", ".join(["%s"] * len(colunas_valores))
            valores = list(colunas_valores.values())
            cur.execute(f"INSERT INTO relatorios_servico ({colunas}) VALUES ({placeholders}) RETURNING id;", valores)
            novo_id = cur.fetchone()["id"]
        conn.commit()
        cur.close()
        conn.close()
        return novo_id, None
    except Exception as e:
        conn.close()
        return None, str(e)

def _normalizar_nome_efetivo(nome, unidade):
    """Se o nome salvo no relatório não bater mais exatamente com a lista atual
    do EFETIVO (ex.: o posto de alguém mudou, como o Rafael hoje), acha o nome
    mais parecido em vez de travar a seleção ou perder o vínculo."""
    lista = EFETIVO.get(unidade, [])
    if nome in lista:
        return nome
    proximos = difflib.get_close_matches(nome, lista, n=1, cutoff=0.6)
    return proximos[0] if proximos else (lista[0] if lista else nome)

def carregar_registro_na_sessao(registro_aberto):
    """Carrega todos os dados de um relatório 'Em Andamento' na sessão atual,
    reutilizada tanto pelo painel 'GU Serviço' quanto pela retomada manual."""
    st.session_state["relatorio_id_atual"] = registro_aberto["id"]
    st.session_state["patrulhamento_terrestre_list"] = json.loads(registro_aberto.get("patrulhamento_terrestre") or "[]")
    st.session_state["patrulhamento_fluvial_list"] = json.loads(registro_aberto.get("patrulhamento_fluvial") or "[]")
    st.session_state["viatura_prefixo"] = registro_aberto.get("viatura_prefixo") or ""
    st.session_state["km_inicial_input"] = registro_aberto.get("km_inicial") or 0
    # Importante: se o relatório já está "Finalizado", o checkbox "Encerrar o
    # serviço agora" precisa vir marcado — senão o campo KM Final some da tela
    # (fica desabilitado) e some do session_state, e quem só abriu pra
    # visualizar/imprimir corre o risco de "perder" o KM Final salvo e ter que
    # redigitar um valor — o que pode mascarar o KM Rodado real se salvar de novo.
    if registro_aberto.get("status") == "Finalizado":
        st.session_state["encerrar_servico_check"] = True
        st.session_state["km_final_input"] = registro_aberto.get("km_final") or 0
    else:
        st.session_state["encerrar_servico_check"] = False
    if registro_aberto.get("data_inicial"):
        st.session_state["data_ini_sel"] = registro_aberto.get("data_inicial")
    if registro_aberto.get("data_final"):
        st.session_state["data_fim_sel"] = registro_aberto.get("data_final")
    _unidade_reg = registro_aberto.get("unidade") or ""
    st.session_state["comandante_sel"] = _normalizar_nome_efetivo(registro_aberto.get("comandante") or "", _unidade_reg)
    st.session_state["motorista_sel"] = _normalizar_nome_efetivo(registro_aberto.get("motorista") or "", _unidade_reg)
    try:
        st.session_state["capturas_animais_list"] = json.loads(registro_aberto.get("capturas_animais") or "[]")
    except Exception:
        st.session_state["capturas_animais_list"] = []
    try:
        st.session_state["apreensoes_list"] = json.loads(registro_aberto.get("apreensoes") or "[]")
    except Exception:
        st.session_state["apreensoes_list"] = []
    try:
        st.session_state["prolepse_list"] = json.loads(registro_aberto.get("visitas_prolepse") or "[]")
    except Exception:
        st.session_state["prolepse_list"] = []
    try:
        st.session_state["armamento_carregado"] = json.loads(registro_aberto.get("armamento_municao") or "[]") or None
    except Exception:
        st.session_state["armamento_carregado"] = None
    st.session_state["alteracoes_servico_input"] = registro_aberto.get("alteracoes_servico") or ""
    _ordens_citadas_str = registro_aberto.get("ordens_servico_citadas") or ""
    st.session_state["ordens_citadas_sel"] = [v.strip() for v in _ordens_citadas_str.split(";") if v.strip()]
    st.session_state["guarnicao_carregada_key"] = f"{registro_aberto.get('unidade')}|{registro_aberto.get('comandante')}"
    # Avisa explicitamente que acabamos de carregar um relatório de propósito —
    # isso impede a proteção "troca de guarnição" (mais abaixo) de disparar por
    # coincidência de texto (ex.: nome do comandante mudou de posto no efetivo)
    # e apagar por engano tudo que acabou de ser carregado.
    st.session_state["_acabou_de_carregar_relatorio"] = True

def montar_dados_parciais(unidade, equipe, finalidade, comandante, motorista, data_ini, data_fim,
                           viatura_p, km_ini):
    """Monta um dicionário parcial (dados dos Blocos 01-02 + listas já registradas)
    para o salvamento automático a cada abordagem/captura. Como é um UPDATE parcial,
    os demais campos (apreensões, armamento, alterações etc.) não são sobrescritos."""
    return {
        "status": "Em Andamento",
        "unidade": unidade,
        "equipe": equipe,
        "finalidade": finalidade,
        "comandante": comandante,
        "motorista": motorista,
        "data_inicial": data_ini,
        "data_final": data_fim,
        "viatura_prefixo": viatura_p,
        "km_inicial": km_ini,
        "patrulhamento_terrestre": json.dumps(st.session_state.get("patrulhamento_terrestre_list", []), ensure_ascii=False),
        "patrulhamento_fluvial": json.dumps(st.session_state.get("patrulhamento_fluvial_list", []), ensure_ascii=False),
        "capturas_animais": json.dumps(st.session_state.get("capturas_animais_list", []), ensure_ascii=False),
        "apreensoes": json.dumps(st.session_state.get("apreensoes_list", []), ensure_ascii=False),
        "visitas_prolepse": json.dumps(st.session_state.get("prolepse_list", []), ensure_ascii=False),
    }

# Busca automática do próximo número sequencial para controle e busca
proximo_numero = 1
conn_controle = init_connection()
if conn_controle:
    try:
        cur_c = conn_controle.cursor()
        cur_c.execute("SELECT COALESCE(MAX(id), 0) + 1 as proximo FROM relatorios_servico;")
        proximo_numero = cur_c.fetchone()['proximo']
        cur_c.close()
        conn_controle.close()
    except:
        proximo_numero = 1

# Inicialização de travas de segurança
if "relatorio_enviado" not in st.session_state:
    st.session_state["relatorio_enviado"] = False

# Dados do Efetivo Real por Extenso
EFETIVO = {
    "2º Pel Miranda": [
        "1º Tenente PM Gesner Batista Ramos",
        "Subtenente PM Luiz Carlos Cavalieri Silva",
        "1º Sargento PM João Vaz",
        "1º Sargento PM Ronaldo da Silva",
        "1º Sargento PM Rafael Bucinsky Fontes",
        "3º Sargento PM Augusto Graça",
        "3º Sargento PM Macsuel Vilalba Santana",
        "Cabo PM Edmar Falcão Santana"
    ],
    "GPM Barra": [
        "3º Sargento PM Luiz Alberto Antonieto",
        "3º Sargento PM Madson Acosta Flores",
        "3º Sargento PM Diego Aguilera Romeiro",
        "Cabo PM Luiz Henrique da Silva Ferreira",
        "Cabo PM Thiago David Mareco de Souza"
    ]
}

# Matrícula de cada policial, usada para preencher automaticamente o campo
# ao selecionar o Comandante da Guarnição. Preencha os que estiverem em branco.
MATRICULAS = {
    "1º Tenente PM Gesner Batista Ramos": "65089021",
    "Subtenente PM Luiz Carlos Cavalieri Silva": "53859021",
    "1º Sargento PM João Vaz": "31702021",
    "1º Sargento PM Ronaldo da Silva": "92954021",
    "1º Sargento PM Rafael Bucinsky Fontes": "95943021",
    "3º Sargento PM Augusto Graça": "45808021",
    "3º Sargento PM Macsuel Vilalba Santana": "117436021",
    "3º Sargento PM Madson Acosta Flores": "25849021",
    "Cabo PM Edmar Falcão Santana": "20305021",
    "3º Sargento PM Luiz Alberto Antonieto": "26999021",
    "3º Sargento PM Diego Aguilera Romeiro": "31260021",
    "Cabo PM Luiz Henrique da Silva Ferreira": "426855021",
    "Cabo PM Thiago David Mareco de Souza": "425383021"
}

# Nome de guerra de cada policial — usado para nomear a equipe automaticamente
# como "Equipe <Nome de Guerra do Comandante>", no padrão militar (GU é do
# comandante). ⚠️ "3º Sargento PM Augusto Graça" não veio na lista que você
# passou — usei "Graça" como um palpite temporário; me avise o nome de guerra
# certo dele para eu corrigir.
NOME_GUERRA = {
    "1º Tenente PM Gesner Batista Ramos": "1º Ten PM Batista",
    "Subtenente PM Luiz Carlos Cavalieri Silva": "ST PM Luiz Carlos",
    "1º Sargento PM João Vaz": "1º Sgt PM João Vaz",
    "1º Sargento PM Ronaldo da Silva": "1º Sgt PM Ronaldo",
    "1º Sargento PM Rafael Bucinsky Fontes": "1º Sgt PM Rafael",
    "3º Sargento PM Augusto Graça": "3º Sgt PM Graça",  # ⚠️ palpite — confirmar nome de guerra
    "3º Sargento PM Macsuel Vilalba Santana": "3º Sgt PM Macsuel",
    "Cabo PM Edmar Falcão Santana": "Cb PM Falcão",
    "3º Sargento PM Luiz Alberto Antonieto": "3º Sgt PM Antonieto",
    "3º Sargento PM Madson Acosta Flores": "3º Sgt PM Madson",
    "3º Sargento PM Diego Aguilera Romeiro": "3º Sgt PM Diego",
    "Cabo PM Luiz Henrique da Silva Ferreira": "Cb PM Luiz Henrique",
    "Cabo PM Thiago David Mareco de Souza": "Cb PM Mareco"
}

# Placas/prefixos de todas as viaturas e embarcações da unidade.
# SUBSTITUA pela lista real assim que você tiver as placas em mãos.
VIATURAS_PLACAS = [
    "REY8G14",
    "RWE6B39",
    "SMK2J66",
    "QAB4447",
    "PBU6375",
    "NRZ4091"
]

# Catálogo de armamento/munição da unidade: código -> nome/descrição.
# Usado na Cautela (item 08) para preencher o nome sozinho a partir do código.
CATALOGO_ARMAMENTO = {
    "31471": "CARABINA MARCA IMBEL MODELO IMBEL IA2 CALIBRE 5,56 Nº SERIE JFA07424",
    "21532": "ESPINGARDA MARCA CBC MODELO CBC MILITARY 3.0 RT TACT 12/19\" CALIBRE 12 Nº SERIE KPB4167550",
    "1626": "FUZIL MARCA FABRIQUE MODELO M964 CALIBRE 7,62 Nº SERIE 26625",
    "17390": "MUNIÇÃO MARCA CBC CALIBRE .40",
    "17392": "MUNIÇÃO MARCA CBC CALIBRE 12",
    "17391": "MUNIÇÃO MARCA CBC CALIBRE 38",
    "17393": "MUNIÇÃO MARCA CBC CALIBRE 5,56",
    "17394": "MUNIÇÃO MARCA CBC CALIBRE 7,62",
    "3736": "PISTOLA MARCA IMBEL MODELO MD7 CALIBRE .40 Nº SERIE EQA01986"
}

# ==========================================
# FUNÇÕES AUXILIARES: DITADO POR VOZ (TRANSCRIÇÃO LOCAL, SEM API PAGA)
# ==========================================
import tempfile
try:
    from faster_whisper import WhisperModel
    VOZ_DISPONIVEL = True
except ImportError:
    VOZ_DISPONIVEL = False

@st.cache_resource
def carregar_modelo_voz():
    """Carrega o modelo de transcrição de voz (Whisper local). Roda 100% offline
    depois do primeiro uso (que baixa o modelo, ~250MB, uma única vez)."""
    return WhisperModel("small", device="cpu", compute_type="int8")

def transcrever_audio(audio_bytes):
    """Transcreve um áudio (bytes) para texto em português."""
    modelo = carregar_modelo_voz()
    with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as tmp:
        tmp.write(audio_bytes)
        caminho_tmp = tmp.name
    segments, _ = modelo.transcribe(caminho_tmp, language="pt")
    return " ".join(seg.text.strip() for seg in segments).strip()

def campo_texto_com_voz(label, key, altura=100, placeholder=None):
    """Campo de texto com botão de ditado por voz ao lado (🎙️). Use no lugar de
    st.text_area(label, key=key) em qualquer campo de relato/observação."""
    col_txt, col_mic = st.columns([9, 1])
    if False:  # comando de voz desativado temporariamente (nunca funcionou direito)
        with col_mic:
            with st.popover("🎙️", use_container_width=True):
                st.caption("Grave sua fala — o texto é transcrito automaticamente.")
                geracao = st.session_state.get(f"audio_geracao_{key}", 0)
                audio = st.audio_input("Gravar", key=f"audio_{key}_{geracao}", label_visibility="collapsed")
                if audio is not None:
                    with st.spinner("Transcrevendo..."):
                        texto_transcrito = transcrever_audio(audio.getvalue())
                    if texto_transcrito:
                        texto_atual = st.session_state.get(key, "")
                        st.session_state[key] = (texto_atual + " " + texto_transcrito).strip()
                    st.session_state[f"audio_geracao_{key}"] = geracao + 1
                    if texto_transcrito:
                        st.success("Transcrito! Confira o texto ao lado.")
                    else:
                        st.warning("Não consegui identificar fala no áudio. Tente novamente.")
                    st.rerun()
    with col_txt:
        valor = st.text_area(label, key=key, height=altura, placeholder=placeholder)
        texto_impressao = html.escape(valor or "—").replace("\n", "<br>")
        st.markdown(
            f'<div class="print-only-text"><strong>{html.escape(label)}:</strong><br>{texto_impressao}</div>',
            unsafe_allow_html=True
        )
    return valor

# ==========================================
# FUNÇÃO AUXILIAR: CAMPO COM MÚLTIPLOS REGISTROS (BOTÃO "+")
# ==========================================
def campo_multiplo(label, session_key, placeholder_prefix=None):
    """Renderiza um campo de texto que pode ser repetido com um botão '+'.
    Retorna a lista de valores preenchidos (não vazios)."""
    if session_key not in st.session_state:
        st.session_state[session_key] = [""]

    if placeholder_prefix is None:
        placeholder_prefix = label

    st.markdown(f"**{label}**")
    remover_idx = None
    geracao = st.session_state.get("doc_geracao", 0)
    for i in range(len(st.session_state[session_key])):
        c1, c2, c3 = st.columns([7, 1, 1])
        with c1:
            st.session_state[session_key][i] = st.text_input(
                label,
                value=st.session_state[session_key][i],
                key=f"{session_key}_{geracao}_{i}",
                label_visibility="collapsed",
                placeholder=f"{placeholder_prefix} nº {i + 1}"
            )
        with c2:
            if i == len(st.session_state[session_key]) - 1:
                if st.button("➕", key=f"add_{session_key}_{geracao}_{i}", help="Adicionar outro registro"):
                    st.session_state[session_key].append("")
                    st.rerun()
        with c3:
            if len(st.session_state[session_key]) > 1:
                if st.button("➖", key=f"rem_{session_key}_{geracao}_{i}", help="Remover este registro"):
                    remover_idx = i

    if remover_idx is not None:
        st.session_state[session_key].pop(remover_idx)
        st.rerun()

    return [v.strip() for v in st.session_state[session_key] if v.strip() != ""]

# Criando as abas principais na interface do usuário
aba_policial, aba_fiscalizacao, aba_ordens, aba_adm = st.tabs(
    ["Formulário de Serviço", "📋 Relatório de Fiscalização", "📑 Ordens de Operação", "Painel Estratégico (Adm)"]
)

# ------------------------------------------
# VISÃO 3: ORDENS DE OPERAÇÃO → ORDENS DE SERVIÇO
# ------------------------------------------
with aba_ordens:
    st.markdown("# 📑 ORDENS DE OPERAÇÃO E ORDENS DE SERVIÇO")
    st.caption("Anexe o PDF da CI/Ordem de Operação do Batalhão. O sistema sugere os campos — você confere e confirma antes de salvar.")

    if not PDF_LEITURA_DISPONIVEL:
        st.warning("⚠️ Leitura de PDF indisponível neste ambiente — adicione `pypdf` ao requirements.txt e reimplante o app. Você ainda pode cadastrar a Ordem preenchendo os campos manualmente abaixo.")

    if "sugestao_ordem" not in st.session_state:
        st.session_state["sugestao_ordem"] = None

    arquivo_ordem_pdf = st.file_uploader("Anexar PDF da CI / Ordem de Operação", type=["pdf"], key="upload_ordem_pdf")
    if arquivo_ordem_pdf is not None and PDF_LEITURA_DISPONIVEL:
        if st.button("🔎 Ler PDF e Sugerir Campos", use_container_width=True):
            try:
                with st.spinner("Lendo o PDF... (se for documento assinado digitalmente, pode levar alguns segundos a mais por usar OCR)"):
                    texto_extraido, metodo_leitura = extrair_texto_pdf(arquivo_ordem_pdf)
                sugestao_nova = extrair_dados_ordem_operacao(texto_extraido)
                st.session_state["sugestao_ordem"] = sugestao_nova

                # Empurra a sugestão direto para as chaves dos widgets abaixo —
                # só usar "value=" não basta aqui, porque essas chaves já existem
                # no session_state desde o primeiro carregamento da tela, e nesse
                # caso o Streamlit ignora "value=" e mantém o que já estava salvo.
                st.session_state["ordem_numero_input"] = sugestao_nova.get("numero") or ""
                if sugestao_nova.get("data_inicio"):
                    st.session_state["ordem_data_inicio_input"] = sugestao_nova["data_inicio"]
                if sugestao_nova.get("data_fim"):
                    st.session_state["ordem_data_fim_input"] = sugestao_nova["data_fim"]
                st.session_state["ordem_finalidade_input"] = sugestao_nova.get("finalidade") or ""
                st.session_state["ordem_cadg_kobo_input"] = sugestao_nova.get("cadg_kobo") or ""

                if not any(sugestao_nova.values()):
                    st.session_state["_msg_leitura_ordem"] = ("warning", "Não consegui identificar nenhum campo automaticamente — talvez o PDF tenha um formato diferente do modelo que vimos. Preencha manualmente abaixo.")
                else:
                    _msg = "Sugestão gerada — confira e ajuste os campos abaixo antes de confirmar."
                    if metodo_leitura == "ocr":
                        _msg = "📸 PDF sem texto selecionável (comum em documento assinado digitalmente) — usei OCR. " + _msg
                    st.session_state["_msg_leitura_ordem"] = ("success", _msg)
                st.rerun()
            except Exception as e:
                st.error(f"Não consegui ler esse PDF: {e}")

    if st.session_state.get("_msg_leitura_ordem"):
        _tipo_msg, _texto_msg = st.session_state.pop("_msg_leitura_ordem")
        getattr(st, _tipo_msg)(_texto_msg)


    sugestao = st.session_state.get("sugestao_ordem") or {"numero": "", "data_inicio": None, "data_fim": None, "finalidade": "", "cadg_kobo": ""}

    st.markdown("### Confirme os dados da Ordem")
    col_ord1, col_ord2, col_ord3 = st.columns(3)
    with col_ord1:
        numero_ordem = st.text_input("Número da Ordem de Operação", value=sugestao.get("numero") or "", key="ordem_numero_input")
    with col_ord2:
        data_inicio_ordem = st.date_input("Data de Início", value=sugestao.get("data_inicio") or datetime.now().date(), key="ordem_data_inicio_input")
    with col_ord3:
        data_fim_ordem = st.date_input("Data Término", value=sugestao.get("data_fim") or datetime.now().date(), key="ordem_data_fim_input")

    def _altura_dinamica(texto, minimo=100, maximo=500):
        """Estima uma altura de caixa de texto proporcional ao conteúdo (o
        Streamlit não tem 'altura automática' nativa, então aproximamos pela
        quantidade de caracteres/linhas)."""
        linhas_estimadas = max(1, len(texto or "") // 90) + (texto or "").count("\n") + 2
        return int(min(maximo, max(minimo, 24 * linhas_estimadas)))

    finalidade_ordem = st.text_area(
        "A equipe deverá realizar... (resumo da finalidade — confira/edite o que o PDF sugeriu)",
        value=sugestao.get("finalidade") or "", height=_altura_dinamica(sugestao.get("finalidade")), key="ordem_finalidade_input"
    )
    cadg_kobo_ordem = st.text_area(
        "Instruções de CADG / KoboToolbox específicas desta Ordem (opcional — se vazio, usa o texto padrão)",
        value=sugestao.get("cadg_kobo") or "", height=_altura_dinamica(sugestao.get("cadg_kobo")), key="ordem_cadg_kobo_input"
    )

    if st.button("✅ Confirmar e Criar Ordem de Serviço", type="primary", use_container_width=True):
        if not numero_ordem or not finalidade_ordem:
            st.error("Preencha ao menos o número da Ordem e a finalidade antes de confirmar.")
        else:
            dados_ordem = {
                "numero": numero_ordem, "data_inicio": data_inicio_ordem, "data_fim": data_fim_ordem,
                "finalidade": finalidade_ordem, "cadg_kobo": cadg_kobo_ordem,
                "arquivo_pdf": arquivo_ordem_pdf.getvalue() if arquivo_ordem_pdf is not None else None,
                "arquivo_pdf_nome": arquivo_ordem_pdf.name if arquivo_ordem_pdf is not None else None,
            }
            novo_id_ordem, erro_ordem = salvar_ordem_operacao(dados_ordem, st.session_state.get("usuario_conectado", ""))
            if erro_ordem:
                st.error(f"Falha ao salvar: {erro_ordem}")
            else:
                st.success(f"✅ Ordem de Serviço criada a partir da Operação {numero_ordem}!")
                st.session_state["sugestao_ordem"] = None
                for _k in ("ordem_numero_input", "ordem_finalidade_input", "ordem_cadg_kobo_input",
                           "ordem_data_inicio_input", "ordem_data_fim_input"):
                    if _k in st.session_state:
                        del st.session_state[_k]
                time.sleep(1)
                st.rerun()

    st.divider()
    st.markdown("### Ordens de Operação Cadastradas")
    todas_ordens = listar_ordens_operacao()
    hoje_ordens = datetime.now().date()
    if not todas_ordens:
        st.info("Nenhuma Ordem de Operação cadastrada ainda.")
    for ordem in todas_ordens:
        _vencida = ordem.get("data_fim") and ordem["data_fim"] < hoje_ordens
        _encerrada_manual = ordem.get("status") == "Encerrada Manualmente"
        if _encerrada_manual:
            _cor, _status_txt = "#888", f"🔒 Encerrada manualmente por {ordem.get('encerrada_por','')}"
        elif _vencida:
            _cor, _status_txt = "#ff4b4b", "🔴 Vencida (data final já passou)"
        else:
            _cor, _status_txt = "#2ecc71", "🟢 Vigente"

        with st.container(border=True):
            col_o1, col_o2 = st.columns([5, 1])
            with col_o1:
                st.markdown(f"**{ordem.get('numero','')}** — {ordem.get('data_inicio','')} a {ordem.get('data_fim','')}")
                st.markdown(f"<span style='color:{_cor}; font-weight:bold;'>{_status_txt}</span>", unsafe_allow_html=True)
                with st.expander("Ver Ordem de Serviço completa"):
                    st.write(montar_texto_ordem_servico(ordem))
            with col_o2:
                if ordem.get("arquivo_pdf"):
                    st.download_button(
                        "📄 PDF Original", data=bytes(ordem["arquivo_pdf"]),
                        file_name=ordem.get("arquivo_pdf_nome") or f"Ordem_{ordem.get('numero','')}.pdf",
                        mime="application/pdf", key=f"baixar_pdf_ordem_{ordem['id']}", use_container_width=True
                    )
                if not _encerrada_manual and not _vencida:
                    if st.button("🔒 Encerrar Antecipadamente", key=f"encerrar_ordem_{ordem['id']}", use_container_width=True):
                        encerrar_ordem_manualmente(ordem["id"], st.session_state.get("usuario_conectado", ""))
                        st.rerun()

# ------------------------------------------
# VISÃO 1: FORMULÁRIO DE SERVIÇO (POLICIAL)
# ------------------------------------------
with aba_policial:
    # Cabeçalho com numeração dinâmica integrada conforme o desenho "Nº"
    col_tit1, col_tit2 = st.columns([4, 1])
    with col_tit1:
        col_logo, col_texto_tit = st.columns([1, 6])
        with col_logo:
            exibir_logo("logo_pantanal_folhas.png", 70)
        with col_texto_tit:
            st.markdown("# RELATÓRIO DE SERVIÇO DIÁRIO - RSD")
            st.caption(f"Unidade: **{st.session_state['unidade_operacional']}** — Preencha os campos operacionais da guarnição abaixo.")
    with col_tit2:
        if st.session_state["relatorio_id_atual"]:
            st.metric("Nº DO RELATÓRIO", f"{st.session_state['relatorio_id_atual']:04d}")
        else:
            st.metric("PRÓXIMO Nº (se novo)", f"{proximo_numero:04d}")

        with st.popover("🔍 Buscar Relatório", use_container_width=True):
            ajuda("Busque por Nº do relatório ou nome do comandante — inclui relatórios já concluídos.")
            termo_busca = st.text_input("Buscar", placeholder="Ex: 0005 ou Madson", key="termo_busca_relatorio", label_visibility="collapsed")
            if termo_busca:
                resultados_busca = buscar_relatorios(st.session_state["unidade_operacional"], termo_busca)
                if resultados_busca:
                    for reg in resultados_busca:
                        st.markdown(f"**Nº {reg['id']:04d}** — {reg.get('comandante','')} — *{reg.get('status','')}*")
                        if reg.get("status") == "Finalizado":
                            ajuda("🔒 Relatório finalizado — protegido contra edição. Só o administrador pode liberar.")
                            senha_desbloqueio = st.text_input(
                                "Senha do administrador para editar", type="password",
                                key=f"senha_desbloqueio_{reg['id']}", label_visibility="collapsed",
                                placeholder="Senha do administrador para editar este relatório"
                            )
                            if st.button("🔓 Desbloquear e Editar", key=f"desbloquear_{reg['id']}", use_container_width=True):
                                if senha_desbloqueio == USUARIOS_PERMITIDOS.get("admin"):
                                    carregar_registro_na_sessao(reg)
                                    registrar_auditoria(reg['id'], "servico", "Liberação de Edição (Admin)", st.session_state.get("usuario_conectado", ""))
                                    st.rerun()
                                else:
                                    st.error("Senha de administrador incorreta. Edição não liberada.")
                        else:
                            if st.button("👁️ Visualizar / Editar", key=f"busca_ver_{reg['id']}", use_container_width=True):
                                carregar_registro_na_sessao(reg)
                                st.rerun()
                        st.divider()
                else:
                    st.info("Nenhum relatório encontrado.")

    # --- PAINEL "GU SERVIÇO" — guarnições com serviço em andamento nesta unidade ---
    # Fica visível assim que a página abre, sem precisar escolher comandante antes.
    # Isolado por unidade_operacional: uma unidade não vê os serviços da outra.
    if st.session_state["relatorio_id_atual"] is None:
        servicos_abertos = listar_relatorios_em_andamento(st.session_state["unidade_operacional"])
        if servicos_abertos:
            st.markdown("##### 🚔 Guarnições com serviço em andamento nesta unidade")
            cols_gu = st.columns(min(len(servicos_abertos), 4))
            for i, reg in enumerate(servicos_abertos):
                try:
                    dias_corridos = (datetime.now().date() - reg["data_inicial"]).days + 1
                except Exception:
                    dias_corridos = 1
                dia_exibido = min(max(dias_corridos, 1), 5)
                nome_curto = (reg.get("comandante") or "").replace("º Sargento PM", "º Sgt").replace("º Tenente PM", "º Ten")
                with cols_gu[i % len(cols_gu)]:
                    if st.button(f"🚔 GU SERVIÇO\nNº {reg['id']:04d} — {nome_curto}\nDia {dia_exibido}/5", key=f"gu_servico_{reg['id']}", use_container_width=True):
                        carregar_registro_na_sessao(reg)
                        st.rerun()
    else:
        st.success(f"🚔 GU em serviço — Nº {st.session_state['relatorio_id_atual']:04d} — Dia em andamento.")

    st.divider()

    st.markdown("### 01 - DADOS DE CONTROLE")
    
    # A equipe agora é derivada automaticamente do Comandante escolhido logo
    # abaixo (padrão militar: a guarnição é do comandante) — ver equipe_sel
    # calculado após a seleção do Comandante da Guarnição.
    
    # Grid de 3 colunas com a data inicial movida para o centro conforme a indicação
    col_u1, col_u2, col_u3 = st.columns(3)
    with col_u1:
        unidade_sel = st.session_state["unidade_operacional"]
        st.text_input("Unidade / Seção", value=unidade_sel, disabled=True)
        # Campo "Finalidade do Serviço" removido da tela: ficou redundante depois que
        # a finalidade passou a ser registrada por atividade nos itens 04 e 05.
        # Mantido nos bastidores só para não quebrar a coluna "finalidade" do banco.
        finalidade_sel = "Patrulhamento Ambiental"
    with col_u2:
        # A Data Inicial assumiu totalmente o lugar central vago
        data_ini_sel = st.date_input("Data Inicial do Serviço", value=datetime.now(), key="data_ini_sel")
    with col_u3:
        data_fim_sel = st.date_input("Data Final do Serviço (Jornada 5 dias)", value=data_ini_sel + timedelta(days=5), key="data_fim_sel")

    _ordens_disponiveis = [
        o for o in listar_ordens_operacao()
        if o.get("status") == "Vigente" and o.get("data_fim") and o["data_fim"] >= datetime.now().date()
    ]
    _opcoes_ordens_sel = [o["numero"] for o in _ordens_disponiveis]
    ordens_citadas_sel = st.multiselect(
        "📑 Ordem(ns) de Serviço vigente(s) aplicável(is) a este relatório",
        options=_opcoes_ordens_sel, key="ordens_citadas_sel",
        help="Vem da aba 'Ordens de Operação'. Se a Operação que você precisa não aparece aqui, cadastre-a primeiro naquela aba."
    )

    st.markdown("#### Guarnição de Serviço")

    # --- APLICA O RESULTADO DA VOZ, SE JÁ CONFIRMADO, ANTES DOS CAMPOS ABAIXO ---
    if st.session_state.get("voz_guarnicao_aplicar"):
        _res = st.session_state.pop("voz_guarnicao_aplicar")
        if _res.get("comandante"):
            st.session_state["comandante_sel"] = _res["comandante"]
        if _res.get("motorista"):
            st.session_state["motorista_sel"] = _res["motorista"]
        if _res.get("viatura"):
            st.session_state["viatura_prefixo"] = _res["viatura"]
        if _res.get("km_inicial") is not None:
            st.session_state["km_inicial_input"] = _res["km_inicial"]

    col_g1, col_g2, col_g3 = st.columns(3)
    with col_g1:
        comandante_sel = st.selectbox("Comandante da Guarnição", EFETIVO[unidade_sel], key="comandante_sel")
    equipe_sel = f"Equipe {NOME_GUERRA.get(comandante_sel, comandante_sel)}"
    with col_g2:
        motorista_sel = st.selectbox("Motorista / Tripulante", EFETIVO[unidade_sel], key="motorista_sel")
    with col_g3:
        matricula_comandante = st.text_input(
            "Matrícula do Comandante",
            value=MATRICULAS.get(comandante_sel, ""),
            placeholder="Ex: 123456-7",
            key=f"matricula_{comandante_sel}"
        )

    # --- PREENCHER GUARNIÇÃO INTEIRA POR VOZ, DE UMA VEZ SÓ ---
    if VOZ_DISPONIVEL:
        with st.popover("🎙️ Preencher Guarnição por Voz", use_container_width=True):
            st.caption("Diga tudo de uma vez, por exemplo: *\"Comandante Sargento Madson, motorista CB Mareco, "
                       "viatura RWE6B39, KM inicial 12000\"*.")
            geracao_g = st.session_state.get("audio_geracao_guarnicao", 0)
            audio_guarnicao = st.audio_input("Gravar guarnição", key=f"audio_guarnicao_{geracao_g}", label_visibility="collapsed")
            if audio_guarnicao is not None:
                with st.spinner("Transcrevendo e interpretando..."):
                    texto_guarnicao = transcrever_audio(audio_guarnicao.getvalue())
                    resultado_voz = interpretar_guarnicao_por_voz(texto_guarnicao, EFETIVO.get(unidade_sel, [])) if texto_guarnicao else None
                st.session_state["audio_geracao_guarnicao"] = geracao_g + 1
                st.session_state["voz_guarnicao_transcricao"] = texto_guarnicao
                st.session_state["voz_guarnicao_resultado"] = resultado_voz
                st.rerun()

    # --- CONFIRMAÇÃO: mostra o que foi entendido antes de aplicar nos campos ---
    if st.session_state.get("voz_guarnicao_resultado") is not None:
        _r = st.session_state["voz_guarnicao_resultado"]
        with st.container(border=True):
            st.markdown(f"🎙️ **Ouvi:** _\"{st.session_state.get('voz_guarnicao_transcricao', '')}\"_")
            st.write(f"**Comandante:** {_r['comandante'] or '⚠️ não identificado'}")
            st.write(f"**Motorista:** {_r['motorista'] or '⚠️ não identificado'}")
            st.write(f"**Viatura:** {_r['viatura'] or '⚠️ não identificado'}")
            st.write(f"**KM Inicial:** {_r['km_inicial'] if _r['km_inicial'] is not None else '⚠️ não identificado'}")
            col_conf1, col_conf2 = st.columns(2)
            with col_conf1:
                if st.button("✅ Confirmar e Preencher", use_container_width=True, key="confirmar_voz_guarnicao"):
                    st.session_state["voz_guarnicao_aplicar"] = _r
                    st.session_state["voz_guarnicao_resultado"] = None
                    st.rerun()
            with col_conf2:
                if st.button("🔁 Descartar e Tentar de Novo", use_container_width=True, key="descartar_voz_guarnicao"):
                    st.session_state["voz_guarnicao_resultado"] = None
                    st.rerun()

    # --- PROTEÇÃO CONTRA TROCA DE GUARNIÇÃO SEM RESET ---
    # Se o comandante selecionado mudou em relação ao serviço carregado na sessão,
    # zera tudo antes de continuar — impede que um "Salvar" acabe sobrescrevendo
    # (por engano) o relatório de outra equipe/guarnição.
    chave_guarnicao_atual = f"{unidade_sel}|{comandante_sel}"
    if st.session_state.pop("_acabou_de_carregar_relatorio", False):
        # Acabamos de carregar um relatório de propósito (busca, GU em serviço,
        # ou desbloqueio de admin) — não reseta, só sincroniza a chave.
        st.session_state["guarnicao_carregada_key"] = chave_guarnicao_atual
    elif st.session_state.get("guarnicao_carregada_key") != chave_guarnicao_atual:
        st.session_state["relatorio_id_atual"] = None
        st.session_state["patrulhamento_terrestre_list"] = []
        st.session_state["patrulhamento_fluvial_list"] = []
        st.session_state["capturas_animais_list"] = []
        st.session_state["apreensoes_list"] = []
        st.session_state["prolepse_list"] = []
        st.session_state["armamento_carregado"] = None
        st.session_state["cadg_list"] = [""]
        st.session_state["auto_infracao_list"] = [""]
        st.session_state["termo_constatacao_list"] = [""]
        st.session_state["termo_apreensao_list"] = [""]
        st.session_state["relatorio_fiscalizacao_list"] = [""]
        st.session_state["relatorio_vistoria_list"] = [""]
        st.session_state["doc_geracao"] = st.session_state.get("doc_geracao", 0) + 1
        st.session_state["guarnicao_carregada_key"] = chave_guarnicao_atual

    st.divider()

    # OBS: os blocos 02 a 05 abaixo NÃO ficam mais dentro de um st.form.
    # Isso é necessário porque o botão "➕" do item 04 precisa disparar uma
    # atualização imediata da tela (rerun), o que não é possível com widgets
    # dentro de um st.form (lá só o botão de envio final reage).

    # Bloco 02: Controle de Viaturas
    st.markdown("### 02 - CONTROLE DE VIATURAS / EMBARCAÇÕES")
    col3, col4, col5 = st.columns(3)
    with col3:
        viatura = st.selectbox("Prefixo da Viatura/Embarcação", VIATURAS_PLACAS, key="viatura_prefixo")
    with col4:
        km_inicial = st.number_input("Quilometragem (KM) Inicial", min_value=0, step=1, key="km_inicial_input")
    with col5:
        encerrar_servico = st.session_state.get("encerrar_servico_check", False)
        km_final = st.number_input(
            "Quilometragem (KM) Final",
            min_value=0, step=1, key="km_final_input",
            disabled=not encerrar_servico,
            help=None if encerrar_servico else "Só é exigido no dia de encerramento do serviço (marque a opção abaixo)."
        )

    # --- TROCA DE ÓLEO: fixa por viatura, não por relatório ---
    ultima_troca_km = buscar_troca_oleo(viatura)
    col_oleo1, col_oleo2 = st.columns([4, 1])
    with col_oleo1:
        st.caption(f"🛢️ Última troca de óleo desta viatura: **{ultima_troca_km:.0f} km**")
    with col_oleo2:
        with st.popover("✏️ Editar Troca de Óleo", use_container_width=True):
            ajuda("Atualize somente quando a troca de óleo desta viatura realmente acontecer.")
            novo_km_troca = st.number_input("Novo KM da troca de óleo", min_value=0, step=1, value=int(ultima_troca_km), key=f"novo_km_troca_{viatura}")
            if st.button("💾 Salvar", key=f"salvar_troca_{viatura}", use_container_width=True):
                if salvar_troca_oleo(viatura, novo_km_troca):
                    st.success("Troca de óleo atualizada!")
                    st.rerun()
                else:
                    st.error("Falha ao salvar — sem conexão com o banco.")

    _situacao_oleo, _msg_oleo = status_troca_oleo(km_inicial, ultima_troca_km)
    if _situacao_oleo == "alerta":
        st.warning(f"🛢️ {_msg_oleo}")
    elif _situacao_oleo == "atrasado":
        st.error(f"🛢️ {_msg_oleo}")

    observacao_viatura = campo_texto_com_voz("Observação sobre a Viatura/Embarcação", "observacao_viatura_input", altura=70)

    encerrar_servico = st.checkbox(
        "🏁 Encerrar o serviço agora (habilita o KM Final e finaliza o relatório de 5 dias)",
        key="encerrar_servico_check"
    )

    km_rodado_calc = calcular_km_rodado(km_inicial, km_final) if encerrar_servico else 0
    st.metric("Distância Total Percorrida (KM Rodado)", f"{km_rodado_calc} km" if encerrar_servico else "Disponível no encerramento")

    @st.dialog("Relatório Salvo")
    def modal_relatorio_finalizado(id_relatorio):
        st.success(f"✅ Relatório Nº {id_relatorio:04d} finalizado e gravado com sucesso na nuvem do Pelotão!")
        st.caption("Os dados continuam na tela até você clicar em Concluir — pode imprimir quantas vezes precisar antes disso.")

        if st.session_state.get("_disparar_impressao_servico"):
            components.html("<script>window.parent.print();</script>", height=0)
            st.session_state["_disparar_impressao_servico"] = False

        col_modal1, col_modal2 = st.columns(2)
        with col_modal1:
            if st.button("🖨️ Imprimir / Salvar em PDF", use_container_width=True):
                registrar_auditoria(id_relatorio, "servico", "Impressão", st.session_state.get("usuario_conectado", ""))
                st.session_state["_disparar_impressao_servico"] = True
                st.rerun()
        with col_modal2:
            if st.button("✅ Concluir e Iniciar Novo Relatório", type="primary", use_container_width=True):
                st.session_state["relatorio_id_atual"] = None
                st.session_state["guarnicao_carregada_key"] = None
                st.session_state["patrulhamento_terrestre_list"] = []
                st.session_state["patrulhamento_fluvial_list"] = []
                st.session_state["capturas_animais_list"] = []
                st.session_state["apreensoes_list"] = []
                st.session_state["prolepse_list"] = []
                st.session_state["armamento_carregado"] = None
                st.session_state["relatorio_aguardando_acao"] = None
                st.rerun()

    if st.session_state.get("relatorio_aguardando_acao"):
        modal_relatorio_finalizado(st.session_state["relatorio_aguardando_acao"])

    def autosave_parcial():
        """Salva automaticamente o progresso atual (Em Andamento) no banco.
        Usado ao inserir, editar ou excluir qualquer item das listas 03/04/05."""
        _dados_auto = montar_dados_parciais(
            unidade_sel, equipe_sel, finalidade_sel, comandante_sel, motorista_sel,
            data_ini_sel, data_fim_sel, viatura, km_inicial
        )
        _id_salvo, _erro_auto = salvar_relatorio(_dados_auto, st.session_state["relatorio_id_atual"])
        if not _erro_auto:
            st.session_state["relatorio_id_atual"] = _id_salvo
        return _erro_auto

    if st.button("💾 Salvar Guarnição / Assumir Serviço", use_container_width=True, key="salvar_guarnicao_btn"):
        _dados_guarnicao = montar_dados_parciais(
            unidade_sel, equipe_sel, finalidade_sel, comandante_sel, motorista_sel,
            data_ini_sel, data_fim_sel, viatura, km_inicial
        )
        _id_salvo, _erro_guarnicao = salvar_relatorio(_dados_guarnicao, st.session_state["relatorio_id_atual"])
        if _erro_guarnicao:
            st.error(f"Falha ao salvar a guarnição: {_erro_guarnicao}")
        else:
            st.session_state["relatorio_id_atual"] = _id_salvo
            st.success(f"✅ Guarnição salva! Serviço Nº {_id_salvo:04d} em andamento — já pode fechar o sistema com segurança e retomar depois.")
            st.rerun()

    st.divider()

    # Bloco 03: Captura de Animais
    st.markdown("### 03 - CAPTURA DE ANIMAIS (PANTANAL)")
    if st.session_state.get("carregar_edicao_animal") is not None:
        _idx_edit = st.session_state.pop("carregar_edicao_animal")
        _item_edit = st.session_state["capturas_animais_list"][_idx_edit]
        st.session_state["editando_idx_animal"] = _idx_edit
        st.session_state["tipo_animal"] = _item_edit.get("ANIMAL", "Não se aplica")
        st.session_state["qtd_animal"] = _item_edit.get("QUANTIDADE", 0)
        st.session_state["especie_animal"] = _item_edit.get("ESPÉCIE", "Não se aplica")
        st.session_state["cadg_animal_input"] = _item_edit.get("Nº CADG", "")
        st.session_state["avaliacao_animal"] = _item_edit.get("AVALIAÇÃO", "Não se aplica")
        st.session_state["origem_animal"] = _item_edit.get("ORIGEM", "Não se aplica")
        st.session_state["destinacao_animal"] = _item_edit.get("DESTINAÇÃO", "")
        st.session_state["relato_captura"] = _item_edit.get("RELATO", "")
    with st.container(border=True):
        col_anim1, col_anim2, col_anim3 = st.columns(3)
        with col_anim1:
            tipo_animal = st.selectbox("Animal Capturado", ["Não se aplica", "Silvestre", "Doméstico"], key="tipo_animal")
            quantidade_animal = st.number_input("Quantidade de Animais", min_value=0, step=1, key="qtd_animal")
        with col_anim2:
            lista_especies = ["Não se aplica", "Tamanduá", "Quati", "Jacaré", "Onça Pintada", "Onça Parda", "Anta", "Tatu", "Periquito", "seriema", "Tucano", "Papagaio","cachorro", "Gato", "Cavalo", "gado", "Cabra", "Carneiro", "Gavião", "Jaguatirica", "Teiú", "Outro"]
            especie_animal = st.selectbox("Espécie do Animal", lista_especies, key="especie_animal")
            cadg_animal = st.text_input("Nº CADG", placeholder="Ex: 12345", key="cadg_animal_input")
        with col_anim3:
            avaliacao_animal = st.selectbox("Avaliação do Estado do Animal", ["Não se aplica", "Ótima", "Boa", "Ruim"], key="avaliacao_animal")
            lista_origens = ["Não se aplica", "Miranda", "Bodoquena", "BR-262", "MS-339", "Residência de Miranda", "Residência de Bodoquena", "Entregue no Pelotão"]
            origem_animal = st.selectbox("Origem / Local da Captura", lista_origens, key="origem_animal")

        col_anim_txt1, col_anim_txt2 = st.columns(2)
        with col_anim_txt1:
            destinacao_animal = campo_texto_com_voz("Destinação do Animal (Breve relato)", "destinacao_animal", altura=70)
        with col_anim_txt2:
            relato_captura = campo_texto_com_voz("Relato Breve da Captura (Como se deu a ação)", "relato_captura", altura=70)

        editando_animal = st.session_state.get("editando_idx_animal")
        label_animal = "💾 Salvar Edição da Captura" if editando_animal is not None else "➕ Inserir / Salvar Captura de Animal"
        col_btn_a1, col_btn_a2 = st.columns([4, 1])
        with col_btn_a1:
            clicou_salvar_animal = st.button(label_animal, use_container_width=True, key="botao_captura_animal")
        with col_btn_a2:
            if editando_animal is not None and st.button("✖️ Cancelar", use_container_width=True, key="cancelar_edicao_animal"):
                st.session_state["editando_idx_animal"] = None
                st.rerun()

        if clicou_salvar_animal:
            nova_captura = {
                "ANIMAL": tipo_animal, "ESPÉCIE": especie_animal, "QUANTIDADE": quantidade_animal,
                "Nº CADG": cadg_animal, "AVALIAÇÃO": avaliacao_animal, "ORIGEM": origem_animal,
                "DESTINAÇÃO": destinacao_animal, "RELATO": relato_captura
            }
            if editando_animal is not None:
                st.session_state["capturas_animais_list"][editando_animal] = nova_captura
                st.session_state["editando_idx_animal"] = None
            else:
                st.session_state["capturas_animais_list"].append(nova_captura)
            _erro_auto = autosave_parcial()
            if _erro_auto:
                st.warning(f"Captura salva localmente, mas o salvamento automático falhou: {_erro_auto}")
            else:
                st.success("Captura de animal salva automaticamente na nuvem!")
            st.rerun()

    if st.session_state["capturas_animais_list"]:
        st.markdown("###### Capturas registradas")
        for i, item in enumerate(st.session_state["capturas_animais_list"]):
            with st.container(border=True):
                col_row1, col_row2, col_row3 = st.columns([6, 1, 1])
                with col_row1:
                    st.markdown(f"**{item.get('ANIMAL','')} — {item.get('ESPÉCIE','')}** | Qtd: {item.get('QUANTIDADE','')} | Nº CADG: {item.get('Nº CADG','') or '—'} | Avaliação: {item.get('AVALIAÇÃO','')} | Origem: {item.get('ORIGEM','')}")
                    if item.get('DESTINAÇÃO') or item.get('RELATO'):
                        st.caption(f"Destinação: {item.get('DESTINAÇÃO','') or '—'} · Relato: {item.get('RELATO','') or '—'}")
                with col_row2:
                    if st.button("✏️ Editar", key=f"edit_animal_{i}", use_container_width=True):
                        st.session_state["carregar_edicao_animal"] = i
                        st.rerun()
                with col_row3:
                    if st.button("🗑️ Excluir", key=f"del_animal_{i}", use_container_width=True):
                        st.session_state["capturas_animais_list"].pop(i)
                        if st.session_state.get("editando_idx_animal") == i:
                            st.session_state["editando_idx_animal"] = None
                        autosave_parcial()
                        st.rerun()
    st.divider()

        # Bloco 04: Patrulhamento Terrestre
    st.markdown("### 04 - PATRULHAMENTO TERRESTRE / FISCALIZAÇÃO / VISTORIA")
    if st.session_state.get("carregar_edicao_terrestre") is not None:
        _idx_edit = st.session_state.pop("carregar_edicao_terrestre")
        _item_edit = st.session_state["patrulhamento_terrestre_list"][_idx_edit]
        st.session_state["editando_idx_terrestre"] = _idx_edit
        st.session_state["sol_t_input"] = _item_edit.get("SOLICITANTE", "Ministério Público")
        st.session_state["fin_t_input"] = _item_edit.get("FINALIDADE", "NUGEO")
        st.session_state["mun_t_input"] = _item_edit.get("MUNICÍPIO", "Miranda")
        st.session_state["pessoas_t_input"] = _item_edit.get("PESSOAS ABORDADAS", 0)
        st.session_state["dist_t_input"] = _item_edit.get("DISTÂNCIA (KM)", 0.0)
        st.session_state["apreenso_t_input"] = _item_edit.get("APREENSÕES", 0)
        st.session_state["vistorias_t_input"] = _item_edit.get("VISTORIAS", 0)
        st.session_state["fiscalizacoes_t_input"] = _item_edit.get("FISCALIZAÇÕES", 0)
        st.session_state["veiculos_t_input"] = _item_edit.get("VEÍCULOS ABORDADOS", 0)
        st.session_state["cadg_t_input"] = _item_edit.get("Nº CADG", "")
        st.session_state["obs_t_input"] = _item_edit.get("OBSERVAÇÕES", "")
    with st.container(border=True):
        col_t1, col_t2, col_t3 = st.columns(3)
        with col_t1:
            sol_t = st.selectbox("Solicitante", ["Ministério Público", "1ºBPMA", "Imasul", "Outros"], key="sol_t_input")
            pessoas_t = st.number_input("Qtd Pessoas Abordadas", min_value=0, step=1, key="pessoas_t_input")
            vistorias_t = st.number_input("Qtd Vistorias", min_value=0, step=1, key="vistorias_t_input")
            cadg_t = st.text_input("Nº CADG", placeholder="Ex: 12345", key="cadg_t_input")
        with col_t2:
            fin_t = st.selectbox("Finalidade", ["NUGEO", "Fiscalização", "Vistoria", "Verificação de Denúncia", "Patrulhamento Preventivo", "Prolepse"], key="fin_t_input")
            dist_t = st.number_input("Distância Percorrida (KM)", min_value=0.0, step=0.1, key="dist_t_input")
            fiscalizacoes_t = st.number_input("Qtd Fiscalizações", min_value=0, step=1, key="fiscalizacoes_t_input")
        with col_t3:
            mun_t = st.selectbox("Município", ["Miranda", "Bodoquena", "Anastácio", "Aquidauana", "Corumbá", "Outros"], key="mun_t_input")
            apreenso_t = st.number_input("Qtd Apreensões", min_value=0, step=1, key="apreenso_t_input")
            veiculos_t = st.number_input("Qtd Veículos Abordados", min_value=0, step=1, key="veiculos_t_input")
        
        obs_t = campo_texto_com_voz("Observações sobre o Policiamento Terrestre", "obs_t_input")
        
        editando_terr = st.session_state.get("editando_idx_terrestre")
        label_terr = "💾 Salvar Edição da Atividade" if editando_terr is not None else "➕ Inserir / Salvar Atividade Terrestre"
        col_btn_t1, col_btn_t2 = st.columns([4, 1])
        with col_btn_t1:
            clicou_salvar_terr = st.button(label_terr, use_container_width=True, key="botao_atividade_terrestre")
        with col_btn_t2:
            if editando_terr is not None and st.button("✖️ Cancelar", use_container_width=True, key="cancelar_edicao_terr"):
                st.session_state["editando_idx_terrestre"] = None
                st.rerun()

        if clicou_salvar_terr:
            nova_atv_t = {
                "SOLICITANTE": sol_t, "FINALIDADE": fin_t, "MUNICÍPIO": mun_t,
                "PESSOAS ABORDADAS": pessoas_t, "DISTÂNCIA (KM)": dist_t,
                "APREENSÕES": apreenso_t, "VISTORIAS": vistorias_t, "FISCALIZAÇÕES": fiscalizacoes_t,
                "VEÍCULOS ABORDADOS": veiculos_t, "Nº CADG": cadg_t,
                "OBSERVAÇÕES": obs_t
            }
            if editando_terr is not None:
                st.session_state["patrulhamento_terrestre_list"][editando_terr] = nova_atv_t
                st.session_state["editando_idx_terrestre"] = None
            else:
                st.session_state["patrulhamento_terrestre_list"].append(nova_atv_t)
            _erro_auto = autosave_parcial()
            if _erro_auto:
                st.warning(f"Atividade salva localmente, mas o salvamento automático falhou: {_erro_auto}")
            else:
                st.success("Atividade terrestre salva automaticamente na nuvem!")
            st.rerun()

    if st.session_state["patrulhamento_terrestre_list"]:
        st.markdown("###### Atividades registradas")
        for i, item in enumerate(st.session_state["patrulhamento_terrestre_list"]):
            with st.container(border=True):
                col_row1, col_row2, col_row3 = st.columns([6, 1, 1])
                with col_row1:
                    st.markdown(f"**{item.get('FINALIDADE','')}** — {item.get('SOLICITANTE','')} | {item.get('MUNICÍPIO','')} | Abordados: {item.get('PESSOAS ABORDADAS','')} | Veículos: {item.get('VEÍCULOS ABORDADOS',0)} | Apreensões: {item.get('APREENSÕES','')} | Vistorias: {item.get('VISTORIAS','')} | Fiscalizações: {item.get('FISCALIZAÇÕES','')} | {item.get('DISTÂNCIA (KM)','')} km")
                    if item.get('OBSERVAÇÕES') or item.get('Nº CADG'):
                        st.caption(f"Nº CADG: {item.get('Nº CADG','') or '—'} · Obs: {item.get('OBSERVAÇÕES','') or '—'}")
                with col_row2:
                    if st.button("✏️ Editar", key=f"edit_terr_{i}", use_container_width=True):
                        st.session_state["carregar_edicao_terrestre"] = i
                        st.rerun()
                with col_row3:
                    if st.button("🗑️ Excluir", key=f"del_terr_{i}", use_container_width=True):
                        st.session_state["patrulhamento_terrestre_list"].pop(i)
                        if st.session_state.get("editando_idx_terrestre") == i:
                            st.session_state["editando_idx_terrestre"] = None
                        autosave_parcial()
                        st.rerun()
    st.divider()

    # Bloco Prolepse: Visitas a Fazendas
    st.markdown("### PROLEPSE - VISITAS A FAZENDAS")
    if st.session_state.get("carregar_edicao_prolepse") is not None:
        _idx_edit = st.session_state.pop("carregar_edicao_prolepse")
        _item_edit = st.session_state["prolepse_list"][_idx_edit]
        st.session_state["editando_idx_prolepse"] = _idx_edit
        st.session_state["fazenda_prolepse"] = _item_edit.get("FAZENDA", "")
        st.session_state["municipio_prolepse"] = _item_edit.get("MUNICÍPIO", "Miranda")
        st.session_state["nome_proprietario_prolepse"] = _item_edit.get("PROPRIETARIO", "")
        st.session_state["contato_prolepse"] = _item_edit.get("CONTATO", "")
        st.session_state["distancia_prolepse"] = _item_edit.get("DISTÂNCIA MIRANDA (KM)", 0.0)
        st.session_state["coordenadas_prolepse"] = _item_edit.get("COORDENADAS", "")
        st.session_state["observacao_prolepse_input"] = _item_edit.get("OBSERVAÇÃO", "")
    with st.container(border=True):
        col_pl1, col_pl2, col_pl3 = st.columns(3)
        with col_pl1:
            fazenda_prolepse = st.text_input("Fazenda Visitada", key="fazenda_prolepse")
            contato_prolepse = st.text_input("Contato do Entrevistado/Fazenda", key="contato_prolepse")
        with col_pl2:
            municipio_prolepse = st.selectbox("Município", ["Miranda", "Bodoquena", "Anastácio", "Aquidauana", "Corumbá", "Outros"], key="municipio_prolepse")
            distancia_prolepse = st.number_input("Distância de Miranda (KM)", min_value=0.0, step=1.0, key="distancia_prolepse")
        with col_pl3:
            nome_proprietario_prolepse = st.text_input("Nome do Proprietario", key="nome_proprietario_prolepse")
            coordenadas_prolepse = st.text_input("Coordenadas", placeholder="Ex: -20.2417, -56.3789", key="coordenadas_prolepse")

        observacao_prolepse = campo_texto_com_voz("Observação", "observacao_prolepse_input", altura=70)

        editando_pl = st.session_state.get("editando_idx_prolepse")
        label_pl = "💾 Salvar Edição da Fazenda" if editando_pl is not None else "➕ Inserir / Salvar Fazenda"
        col_btn_pl1, col_btn_pl2 = st.columns([4, 1])
        with col_btn_pl1:
            clicou_salvar_pl = st.button(label_pl, use_container_width=True, key="botao_prolepse")
        with col_btn_pl2:
            if editando_pl is not None and st.button("✖️ Cancelar", use_container_width=True, key="cancelar_edicao_pl"):
                st.session_state["editando_idx_prolepse"] = None
                st.rerun()

        if clicou_salvar_pl:
            nova_visita_prolepse = {
                "FAZENDA": fazenda_prolepse, "MUNICÍPIO": municipio_prolepse,
                "PROPRIETARIO": nome_proprietario_prolepse, "CONTATO": contato_prolepse,
                "DISTÂNCIA MIRANDA (KM)": distancia_prolepse, "COORDENADAS": coordenadas_prolepse,
                "OBSERVAÇÃO": observacao_prolepse
            }
            if editando_pl is not None:
                st.session_state["prolepse_list"][editando_pl] = nova_visita_prolepse
                st.session_state["editando_idx_prolepse"] = None
            else:
                st.session_state["prolepse_list"].append(nova_visita_prolepse)
            _erro_auto = autosave_parcial()
            if _erro_auto:
                st.warning(f"Fazenda salva localmente, mas o salvamento automático falhou: {_erro_auto}")
            else:
                st.success("Visita à fazenda salva automaticamente na nuvem!")
            st.rerun()

    if st.session_state["prolepse_list"]:
        st.markdown("###### Fazendas visitadas")
        for i, item in enumerate(st.session_state["prolepse_list"]):
            with st.container(border=True):
                col_row1, col_row2, col_row3 = st.columns([6, 1, 1])
                with col_row1:
                    st.markdown(f"**{item.get('FAZENDA','')}** — {item.get('MUNICÍPIO','')} | Proprietário: {item.get('PROPRIETARIO','')} | Contato: {item.get('CONTATO','') or '—'} | {item.get('DISTÂNCIA MIRANDA (KM)','')} km de Miranda")
                    if item.get('COORDENADAS') or item.get('OBSERVAÇÃO'):
                        st.caption(f"Coordenadas: {item.get('COORDENADAS','') or '—'} · Obs: {item.get('OBSERVAÇÃO','') or '—'}")
                with col_row2:
                    if st.button("✏️ Editar", key=f"edit_pl_{i}", use_container_width=True):
                        st.session_state["carregar_edicao_prolepse"] = i
                        st.rerun()
                with col_row3:
                    if st.button("🗑️ Excluir", key=f"del_pl_{i}", use_container_width=True):
                        st.session_state["prolepse_list"].pop(i)
                        if st.session_state.get("editando_idx_prolepse") == i:
                            st.session_state["editando_idx_prolepse"] = None
                        autosave_parcial()
                        st.rerun()
    st.divider()

    # Bloco 05: Patrulhamento Fluvial
    st.markdown("### 05 - PATRULHAMENTO FLUVIAL")
    if st.session_state.get("carregar_edicao_fluvial") is not None:
        _idx_edit = st.session_state.pop("carregar_edicao_fluvial")
        _item_edit = st.session_state["patrulhamento_fluvial_list"][_idx_edit]
        st.session_state["editando_idx_fluvial"] = _idx_edit
        st.session_state["sol_f_input"] = _item_edit.get("SOLICITANTE", "Ministério Público")
        st.session_state["fin_f_input"] = _item_edit.get("FINALIDADE", "NUGEO")
        st.session_state["mun_f_input"] = _item_edit.get("MUNICÍPIO", "Miranda")
        st.session_state["pescadores_f_input"] = _item_edit.get("PESCADORES ABORDADOS", 0)
        st.session_state["dist_f_input"] = _item_edit.get("DISTÂNCIA", 0.0)
        st.session_state["apreenso_f_input"] = _item_edit.get("APREENSÕES", 0)
        st.session_state["vistorias_f_input"] = _item_edit.get("VISTORIAS", 0)
        st.session_state["embarcacoes_f_input"] = _item_edit.get("EMBARCAÇÕES ABORDADAS", 0)
        st.session_state["cadg_f_input"] = _item_edit.get("Nº CADG", "")
        st.session_state["obs_f_input"] = _item_edit.get("OBSERVAÇÕES", "")
    with st.container(border=True):
        col_f1, col_f2, col_f3 = st.columns(3)
        with col_f1:
            sol_f = st.selectbox("Solicitante", ["Ministério Público", "1ºBPMA", "Imasul", "Outros"], key="sol_f_input")
            pescadores_f = st.number_input("Qtd Pescadores Abordados", min_value=0, step=1, key="pescadores_f_input")
            vistorias_f = st.number_input("Qtd Vistorias (Fluvial)", min_value=0, step=1, key="vistorias_f_input")
        with col_f2:
            fin_f = st.selectbox("Finalidade", ["NUGEO", "Fiscalização", "Vistoria", "Verificação de Denúncia", "Patrulhamento Preventivo", "Prolepse"], key="fin_f_input")
            dist_f = st.number_input("Distância Percorrida (KM/Milhas)", min_value=0.0, step=0.1, key="dist_f_input")
            apreenso_f = st.number_input("Qtd Apreensões (Fluvial)", min_value=0, step=1, key="apreenso_f_input")
        with col_f3:
            mun_f = st.selectbox("Município", ["Miranda", "Bodoquena", "Anastácio", "Aquidauana", "Corumbá", "Outros"], key="mun_f_input")
            embarcacoes_f = st.number_input("Qtd Embarcações Abordadas", min_value=0, step=1, key="embarcacoes_f_input")
            cadg_f = st.text_input("Nº CADG", placeholder="Ex: 12345", key="cadg_f_input")
            
        obs_f = campo_texto_com_voz("Observações sobre o Policiamento Fluvial", "obs_f_input")
        
        editando_fluv = st.session_state.get("editando_idx_fluvial")
        label_fluv = "💾 Salvar Edição da Atividade" if editando_fluv is not None else "➕ Inserir / Salvar Atividade Fluvial"
        col_btn_f1, col_btn_f2 = st.columns([4, 1])
        with col_btn_f1:
            clicou_salvar_fluv = st.button(label_fluv, use_container_width=True, key="botao_atividade_fluvial")
        with col_btn_f2:
            if editando_fluv is not None and st.button("✖️ Cancelar", use_container_width=True, key="cancelar_edicao_fluv"):
                st.session_state["editando_idx_fluvial"] = None
                st.rerun()

        if clicou_salvar_fluv:
            nova_atv_f = {
                "SOLICITANTE": sol_f, "FINALIDADE": fin_f, "MUNICÍPIO": mun_f,
                "EMBARCAÇÕES ABORDADAS": embarcacoes_f, "PESCADORES ABORDADOS": pescadores_f,
                "DISTÂNCIA": dist_f, "APREENSÕES": apreenso_f, "VISTORIAS": vistorias_f,
                "Nº CADG": cadg_f, "OBSERVAÇÕES": obs_f
            }
            if editando_fluv is not None:
                st.session_state["patrulhamento_fluvial_list"][editando_fluv] = nova_atv_f
                st.session_state["editando_idx_fluvial"] = None
            else:
                st.session_state["patrulhamento_fluvial_list"].append(nova_atv_f)
            _erro_auto = autosave_parcial()
            if _erro_auto:
                st.warning(f"Atividade salva localmente, mas o salvamento automático falhou: {_erro_auto}")
            else:
                st.success("Atividade fluvial salva automaticamente na nuvem!")
            st.rerun()

    if st.session_state["patrulhamento_fluvial_list"]:
        st.markdown("###### Atividades registradas")
        for i, item in enumerate(st.session_state["patrulhamento_fluvial_list"]):
            with st.container(border=True):
                col_row1, col_row2, col_row3 = st.columns([6, 1, 1])
                with col_row1:
                    st.markdown(f"**{item.get('FINALIDADE','')}** — {item.get('SOLICITANTE','')} | {item.get('MUNICÍPIO','')} | Embarcações: {item.get('EMBARCAÇÕES ABORDADAS','')} | Pescadores: {item.get('PESCADORES ABORDADOS','')} | Apreensões: {item.get('APREENSÕES','')} | Vistorias: {item.get('VISTORIAS','')} | {item.get('DISTÂNCIA','')} km")
                    if item.get('OBSERVAÇÕES') or item.get('Nº CADG'):
                        st.caption(f"Nº CADG: {item.get('Nº CADG','') or '—'} · Obs: {item.get('OBSERVAÇÕES','') or '—'}")
                with col_row2:
                    if st.button("✏️ Editar", key=f"edit_fluv_{i}", use_container_width=True):
                        st.session_state["carregar_edicao_fluvial"] = i
                        st.rerun()
                with col_row3:
                    if st.button("🗑️ Excluir", key=f"del_fluv_{i}", use_container_width=True):
                        st.session_state["patrulhamento_fluvial_list"].pop(i)
                        if st.session_state.get("editando_idx_fluvial") == i:
                            st.session_state["editando_idx_fluvial"] = None
                        autosave_parcial()
                        st.rerun()
    st.divider()

    # Bloco 06: Apreensões
    st.markdown("### 06 - APREENSÕES")
    if st.session_state.get("carregar_edicao_apreensao") is not None:
        _idx_edit = st.session_state.pop("carregar_edicao_apreensao")
        _item_edit = st.session_state["apreensoes_list"][_idx_edit]
        st.session_state["editando_idx_apreensao"] = _idx_edit
        st.session_state["infra_crime"] = _item_edit.get("INFRAÇÃO/CRIME", "Não se aplica")
        st.session_state["desc_material"] = _item_edit.get("DESCRIÇÃO", "")
        st.session_state["municipio_apreensao"] = _item_edit.get("MUNICÍPIO", "Não se aplica")
        st.session_state["tipo_material"] = _item_edit.get("TIPO MATERIAL", "Não se aplica")
        st.session_state["quantidade_apreendida"] = _item_edit.get("QUANTIDADE", 0.0)
        st.session_state["unidade_medida"] = _item_edit.get("UNIDADE MEDIDA", "Unidades")
        st.session_state["valor_multa"] = _item_edit.get("VALOR MULTA", 0.0)
        st.session_state["relato_apreensao"] = _item_edit.get("RELATO", "")
        st.session_state["cadg_list"] = _item_edit.get("CADG", "").split("; ") if _item_edit.get("CADG") else [""]
        st.session_state["auto_infracao_list"] = _item_edit.get("AUTO INFRAÇÃO", "").split("; ") if _item_edit.get("AUTO INFRAÇÃO") else [""]
        st.session_state["termo_constatacao_list"] = _item_edit.get("TERMO CONSTATAÇÃO", "").split("; ") if _item_edit.get("TERMO CONSTATAÇÃO") else [""]
        st.session_state["termo_apreensao_list"] = _item_edit.get("TERMO APREENSÃO", "").split("; ") if _item_edit.get("TERMO APREENSÃO") else [""]
        st.session_state["relatorio_fiscalizacao_list"] = _item_edit.get("RELATÓRIO FISCALIZAÇÃO", "").split("; ") if _item_edit.get("RELATÓRIO FISCALIZAÇÃO") else [""]
        st.session_state["relatorio_vistoria_list"] = _item_edit.get("RELATÓRIO VISTORIA", "").split("; ") if _item_edit.get("RELATÓRIO VISTORIA") else [""]
        st.session_state["doc_geracao"] = st.session_state.get("doc_geracao", 0) + 1
    with st.container(border=True):
        col_ap1, col_ap2, col_ap3 = st.columns(3)
        with col_ap1:
            lista_crimes = [
                "Não se aplica", "Pesca Ilegal / Predatória", "Caça Ilegal", 
                "Desmatamento / Exploração Florestal", "Transporte Irregular de Madeira/Carvão",
                "Tráfico de Drogas", "Contrabando", "Descaminho", "Outros Crimes/Infrações"
            ]
            infra_crime = st.selectbox("Infração / Crime Constatado", lista_crimes, key="infra_crime")
            desc_material = st.text_input("Descrição Detalhada do Material", placeholder="Ex: 15kg de Pacu, Rede de pesca, 2m³ de aroeira", key="desc_material")
        with col_ap2:
            municipio_apreensao = st.selectbox("Município da Apreensão", ["Não se aplica", "Miranda", "Bodoquena", "Outros"], key="municipio_apreensao")
            tipo_material = st.selectbox("Tipo de Material Apreendido", ["Não se aplica", "Pescado", "Madeira", "Fauna (outros)", "Flora (outros)", "Outros"], key="tipo_material")
        with col_ap3:
            quantidade_apreendida = st.number_input("Quantidade (Em unidade ou medida)", min_value=0.0, step=1.0, key="quantidade_apreendida")
            unidade_medida = st.selectbox("Unidade da Quantidade Acima", ["Unidades", "Kg", "m³", "Litros", "Outra"], key="unidade_medida")
            valor_multa = st.number_input("Valor da Multa Aplicada (R$)", min_value=0.0, step=50.0, key="valor_multa")

        relato_apreensao = campo_texto_com_voz("Breve Relato da Apreensão / Ocorrência", "relato_apreensao")

        st.divider()
        st.markdown("#### Documentos Vinculados a Esta Apreensão")
        ajuda("Use o botão ➕ para registrar mais de um documento do mesmo tipo nesta mesma apreensão.")
        col_doc1, col_doc2 = st.columns(2)
        with col_doc1:
            cadg_valores = campo_multiplo("Nº CADG", "cadg_list")
            termo_constatacao_valores = campo_multiplo("Nº Termo de Constatação", "termo_constatacao_list")
            relatorio_fiscalizacao_valores = campo_multiplo("Nº Relatório de Fiscalização", "relatorio_fiscalizacao_list")
        with col_doc2:
            auto_infracao_valores = campo_multiplo("Nº Auto de Infração", "auto_infracao_list")
            termo_apreensao_valores = campo_multiplo("Nº Termo de Apreensão", "termo_apreensao_list")
            relatorio_vistoria_valores = campo_multiplo("Nº Relatório de Vistoria", "relatorio_vistoria_list")

        st.divider()
        editando_ap = st.session_state.get("editando_idx_apreensao")
        label_ap = "💾 Salvar Edição da Apreensão" if editando_ap is not None else "➕ Inserir / Salvar Apreensão"
        col_btn_ap1, col_btn_ap2 = st.columns([4, 1])
        with col_btn_ap1:
            clicou_salvar_ap = st.button(label_ap, use_container_width=True, key="botao_apreensao")
        with col_btn_ap2:
            if editando_ap is not None and st.button("✖️ Cancelar", use_container_width=True, key="cancelar_edicao_ap"):
                st.session_state["editando_idx_apreensao"] = None
                st.rerun()

        if clicou_salvar_ap:
            nova_apreensao = {
                "INFRAÇÃO/CRIME": infra_crime, "TIPO MATERIAL": tipo_material, "DESCRIÇÃO": desc_material,
                "QUANTIDADE": quantidade_apreendida, "UNIDADE MEDIDA": unidade_medida, "MUNICÍPIO": municipio_apreensao,
                "VALOR MULTA": valor_multa, "RELATO": relato_apreensao,
                "CADG": "; ".join(cadg_valores),
                "AUTO INFRAÇÃO": "; ".join(auto_infracao_valores),
                "TERMO CONSTATAÇÃO": "; ".join(termo_constatacao_valores),
                "TERMO APREENSÃO": "; ".join(termo_apreensao_valores),
                "RELATÓRIO FISCALIZAÇÃO": "; ".join(relatorio_fiscalizacao_valores),
                "RELATÓRIO VISTORIA": "; ".join(relatorio_vistoria_valores),
            }
            if editando_ap is not None:
                st.session_state["apreensoes_list"][editando_ap] = nova_apreensao
                st.session_state["editando_idx_apreensao"] = None
            else:
                st.session_state["apreensoes_list"].append(nova_apreensao)
            # Limpa os documentos para a próxima apreensão não herdar os desta
            st.session_state["cadg_list"] = [""]
            st.session_state["auto_infracao_list"] = [""]
            st.session_state["termo_constatacao_list"] = [""]
            st.session_state["termo_apreensao_list"] = [""]
            st.session_state["relatorio_fiscalizacao_list"] = [""]
            st.session_state["relatorio_vistoria_list"] = [""]
            st.session_state["doc_geracao"] = st.session_state.get("doc_geracao", 0) + 1
            _erro_auto = autosave_parcial()
            if _erro_auto:
                st.warning(f"Apreensão salva localmente, mas o salvamento automático falhou: {_erro_auto}")
            else:
                st.success("Apreensão e documentos vinculados salvos automaticamente na nuvem!")
            st.rerun()

    if st.session_state["apreensoes_list"]:
        st.markdown("###### Apreensões registradas")
        for i, item in enumerate(st.session_state["apreensoes_list"]):
            with st.container(border=True):
                col_row1, col_row2, col_row3 = st.columns([6, 1, 1])
                with col_row1:
                    st.markdown(f"**{item.get('INFRAÇÃO/CRIME','')}** — {item.get('TIPO MATERIAL','')} | {item.get('MUNICÍPIO','')} | Qtd: {item.get('QUANTIDADE','')} {item.get('UNIDADE MEDIDA','')} | Multa: R$ {item.get('VALOR MULTA',0):,.2f}")
                    if item.get('DESCRIÇÃO') or item.get('RELATO'):
                        st.caption(f"Descrição: {item.get('DESCRIÇÃO','') or '—'} · Relato: {item.get('RELATO','') or '—'}")
                    docs_resumo = " · ".join(f"{rotulo}: {item.get(rotulo)}" for rotulo in ["CADG", "AUTO INFRAÇÃO", "TERMO CONSTATAÇÃO", "TERMO APREENSÃO", "RELATÓRIO FISCALIZAÇÃO", "RELATÓRIO VISTORIA"] if item.get(rotulo))
                    if docs_resumo:
                        st.caption(f"📎 {docs_resumo}")
                with col_row2:
                    if st.button("✏️ Editar", key=f"edit_ap_{i}", use_container_width=True):
                        st.session_state["carregar_edicao_apreensao"] = i
                        st.rerun()
                with col_row3:
                    if st.button("🗑️ Excluir", key=f"del_ap_{i}", use_container_width=True):
                        st.session_state["apreensoes_list"].pop(i)
                        if st.session_state.get("editando_idx_apreensao") == i:
                            st.session_state["editando_idx_apreensao"] = None
                        autosave_parcial()
                        st.rerun()

    st.divider()

        # Bloco 07: Estatística (Automatizado)
    st.markdown("### 07 - ESTATÍSTICA CONSOLIDADA DO SERVIÇO")
    
    # Cálculos automáticos baseados nas tabelas salvas pelo agente
    total_pessoas_t = sum(item["PESSOAS ABORDADAS"] for item in st.session_state["patrulhamento_terrestre_list"])
    total_dist_t = sum(item["DISTÂNCIA (KM)"] for item in st.session_state["patrulhamento_terrestre_list"])
    total_apre_t = sum(item["APREENSÕES"] for item in st.session_state["patrulhamento_terrestre_list"])
    total_vist_t = sum(item["VISTORIAS"] for item in st.session_state["patrulhamento_terrestre_list"])
    total_fisc_t = sum(item["FISCALIZAÇÕES"] for item in st.session_state["patrulhamento_terrestre_list"])
    
    total_emb_f = sum(item["EMBARCAÇÕES ABORDADAS"] for item in st.session_state["patrulhamento_fluvial_list"])
    total_pesc_f = sum(item["PESCADORES ABORDADOS"] for item in st.session_state["patrulhamento_fluvial_list"])
    total_dist_f = sum(item["DISTÂNCIA"] for item in st.session_state["patrulhamento_fluvial_list"])
    total_apre_f = sum(item["APREENSÕES"] for item in st.session_state["patrulhamento_fluvial_list"])
    total_vist_f = sum(item["VISTORIAS"] for item in st.session_state["patrulhamento_fluvial_list"])
    total_prolepse = sum(1 for item in st.session_state["patrulhamento_terrestre_list"] if item["FINALIDADE"] == "Prolepse") + \
                      sum(1 for item in st.session_state["patrulhamento_fluvial_list"] if item["FINALIDADE"] == "Prolepse")
    total_veiculos_t = sum(item.get("VEÍCULOS ABORDADOS", 0) for item in st.session_state["patrulhamento_terrestre_list"])

    total_ai = 0
    total_multas_valor = 0.0
    total_pescado_un = 0.0
    total_pescado_kg = 0.0
    total_madeira_un = 0.0
    total_madeira_m3 = 0.0
    for ap in st.session_state["apreensoes_list"]:
        auto_str = ap.get("AUTO INFRAÇÃO", "") or ""
        total_ai += len([v for v in auto_str.split("; ") if v.strip()])
        total_multas_valor += ap.get("VALOR MULTA", 0) or 0
        qtd_ap = ap.get("QUANTIDADE", 0) or 0
        unidade_ap = ap.get("UNIDADE MEDIDA", "")
        if ap.get("TIPO MATERIAL") == "Pescado":
            if unidade_ap == "Kg":
                total_pescado_kg += qtd_ap
            elif unidade_ap == "Unidades":
                total_pescado_un += qtd_ap
        elif ap.get("TIPO MATERIAL") == "Madeira":
            if unidade_ap == "m³":
                total_madeira_m3 += qtd_ap
            elif unidade_ap == "Unidades":
                total_madeira_un += qtd_ap

    # Exibição visual limpa em formato de cards para conferência rápida
    st.info("📊 Os dados abaixo são calculados em tempo real com base nos patrulhamentos e apreensões inseridos acima.")
    c_est1, c_est2, c_est3 = st.columns(3)
    with c_est1:
        st.metric("Total Pessoas Abordadas (Terr.)", total_pessoas_t)
        st.metric("Total Pescadores Abordados", total_pesc_f)
        st.metric("Total Vistorias (Geral)", total_vist_t + total_vist_f)
        st.metric("Total Veículos Abordados", total_veiculos_t)
    with c_est2:
        st.metric("Total Distância Terrestre", f"{total_dist_t:.1f} KM")
        st.metric("Total Distância Fluvial", f"{total_dist_f:.1f} KM/MN")
        st.metric("Total Fiscalizações", total_fisc_t)
        st.metric("Total Autos de Infração (AI)", total_ai)
    with c_est3:
        st.metric("Total Embarcações Abordadas", total_emb_f)
        st.metric("Total Apreensões (Geral)", total_apre_t + total_apre_f)
        st.metric("Total Prolepse", total_prolepse)
        st.metric("Total em Multas Aplicadas", f"R$ {total_multas_valor:,.2f}")

    st.markdown("###### Apreensões por Tipo de Material")
    c_est4, c_est5 = st.columns(2)
    with c_est4:
        st.metric("Total Pescado Apreendido", f"{total_pescado_un:.0f} un / {total_pescado_kg:.1f} Kg")
    with c_est5:
        st.metric("Total Madeira Apreendida", f"{total_madeira_un:.0f} un / {total_madeira_m3:.1f} m³")

    st.divider()


        # Bloco 08: Controle de Armamento e Munição
    st.markdown("### 08 - CONTROLE DE ARMAMENTO E MUNIÇÃO")
    st.write("Utilize a tabela abaixo para gerenciar os armamentos e munições. Clique em '+' no final da tabela para adicionar novas linhas ou selecione e pressione 'Delete' para excluir.")
    
    # Dados padrão baseados na sua imagem para preenchimento inicial automático
    dados_iniciais = st.session_state.get("armamento_carregado") or [
        {"CODIGO": "31471", "NOMENCLATURA": "CARABINA MARCA IMBEL MODELO IMBEL IA2 CALIBRE 5,56 Nº SERIE JFA07424", "QTDE": 1.00},
        {"CODIGO": "21532", "NOMENCLATURA": "ESPINGARDA MARCA CBC MODELO CBC MILITARY 3.0 RT TACT 12/19\" CALIBRE 12 Nº SERIE KPB4167550", "QTDE": 1.00},
        {"CODIGO": "1626", "NOMENCLATURA": "FUZIL MARCA FABRIQUE MODELO M964 CALIBRE 7,62 Nº SERIE 26625", "QTDE": 1.00},
        {"CODIGO": "17390", "NOMENCLATURA": "MUNICAO MARCA CBC CALIBRE .40", "QTDE": 150.00},
        {"CODIGO": "17392", "NOMENCLATURA": "MUNICAO MARCA CBC CALIBRE 12", "QTDE": 50.00},
        {"CODIGO": "17391", "NOMENCLATURA": "MUNICAO MARCA CBC CALIBRE 38", "QTDE": 25.00},
        {"CODIGO": "17393", "NOMENCLATURA": "MUNICAO MARCA CBC CALIBRE 5,56", "QTDE": 100.00},
        {"CODIGO": "17394", "NOMENCLATURA": "MUNICAO MARCA CBC CALIBRE 7,62", "QTDE": 50.00},
        {"CODIGO": "3736", "NOMENCLATURA": "PISTOLA MARCA IMBEL MODELO MD7 CALIBRE .40 Nº SERIE EQA01986", "QTDE": 1.00}
    ]
    
    # Data editor interativo que permite adicionar, editar e remover linhas.
    # A chave muda por guarnição (não por relatório) para não perder o que já
    # está sendo digitado quando o Nº do relatório é criado no meio da sessão,
    # mas ainda assim recarregar do zero ao trocar de guarnição/retomar outra.
    tabela_armamento = st.data_editor(
        dados_iniciais,
        num_rows="dynamic",
        use_container_width=True,
        key=f"armamento_editor_{st.session_state.get('guarnicao_carregada_key', 'novo')}",
        column_config={
            "CODIGO": st.column_config.TextColumn("CÓDIGO"),
            "NOMENCLATURA": st.column_config.TextColumn("NOMENCLATURA/DESCRIÇÃO", width="large"),
            "QTDE": st.column_config.NumberColumn("QTDE", min_value=0.0, format="%.2f")
        }
    )

    observacao_armamento = campo_texto_com_voz("Observação sobre Armamento e Munição", "observacao_armamento_input", altura=70)

    # ================= CAUTELA DE ARMAMENTO E MUNIÇÃO =================
    st.markdown("####  Cautela de Armamento e Munição")
    ajuda("A cautela fica em aberto (visível em todos os relatórios da unidade) até que o material seja devolvido.")

    with st.popover(" Nova Cautela", use_container_width=True):
        st.markdown("**1. Adicione os itens da cautela**")
        col_cau1, col_cau2 = st.columns([2, 1])
        with col_cau1:
            codigo_cautela = st.text_input("Código do item", key="codigo_cautela_input")
        with col_cau2:
            qtd_cautela = st.number_input("Quantidade", min_value=1, step=1, key="qtd_cautela_input")

        nome_item_cautela = CATALOGO_ARMAMENTO.get(codigo_cautela.strip()) if codigo_cautela else None
        if codigo_cautela and not nome_item_cautela:
            st.warning("Código não encontrado no catálogo.")
        elif nome_item_cautela:
            st.caption(f"📦 {nome_item_cautela}")

        if st.button("➕ Adicionar Item à Cautela", use_container_width=True, disabled=not nome_item_cautela):
            st.session_state["cautela_itens_temp"].append({
                "CÓDIGO": codigo_cautela.strip(), "NOME": nome_item_cautela, "QUANTIDADE": qtd_cautela
            })
            st.rerun()

        if st.session_state["cautela_itens_temp"]:
            st.markdown("**Itens já adicionados:**")
            for i, item_cau in enumerate(st.session_state["cautela_itens_temp"]):
                col_ci1, col_ci2 = st.columns([5, 1])
                with col_ci1:
                    st.caption(f"{item_cau['CÓDIGO']} — {item_cau['NOME']} (Qtd: {item_cau['QUANTIDADE']})")
                with col_ci2:
                    if st.button("🗑️", key=f"del_cautela_temp_{i}"):
                        st.session_state["cautela_itens_temp"].pop(i)
                        st.rerun()

            st.divider()
            st.markdown("**2. Dados da cautela**")
            destinatario_cautela = st.text_input("Para quem está cautelando", key="destinatario_cautela_input")
            col_cau3, col_cau4 = st.columns(2)
            with col_cau3:
                data_cautela_val = st.date_input("Data da Cautela", value=datetime.now(), key="data_cautela_input")
            with col_cau4:
                prazo_indefinido = st.checkbox("Prazo indefinido", key="prazo_indefinido_cautela")
                prazo_cautela_val = st.date_input("Prazo para devolução", value=datetime.now(), key="prazo_cautela_input", disabled=prazo_indefinido)

            if st.button("✅ Finalizar Cautela", type="primary", use_container_width=True):
                if not destinatario_cautela:
                    st.error("Informe para quem está cautelando.")
                else:
                    novo_id_cautela, erro_cautela = criar_cautela(
                        st.session_state["unidade_operacional"], destinatario_cautela, data_cautela_val,
                        None if prazo_indefinido else prazo_cautela_val, prazo_indefinido,
                        st.session_state["cautela_itens_temp"]
                    )
                    if erro_cautela:
                        st.error(f"Falha ao salvar a cautela: {erro_cautela}")
                    else:
                        st.session_state["cautela_itens_temp"] = []
                        st.success(f"✅ Cautela Nº {novo_id_cautela:04d} finalizada!")
                        st.rerun()

    # --- CAUTELAS EM ABERTO DESTA UNIDADE ---
    cautelas_abertas = listar_cautelas_abertas(st.session_state["unidade_operacional"])
    if cautelas_abertas:
        st.markdown("##### 📋 Cautelas em Aberto")
        for cautela in cautelas_abertas:
            try:
                itens_cautela = json.loads(cautela.get("itens") or "[]")
            except Exception:
                itens_cautela = []
            with st.container(border=True):
                prazo_txt = "Indefinido" if cautela.get("prazo_indefinido") else str(cautela.get("prazo") or "—")
                st.markdown(f"**Cautela Nº {cautela['id']:04d}** — Para: {cautela.get('destinatario','')} | Data: {cautela.get('data_cautela','')} | Prazo: {prazo_txt}")
                for item_cau in itens_cautela:
                    st.caption(f"• {item_cau.get('CÓDIGO','')} — {item_cau.get('NOME','')} (Qtd: {item_cau.get('QUANTIDADE','')})")

                with st.popover("📦 Entrega de Materiais", use_container_width=True):
                    ajuda("Confirme a devolução do material desta cautela.")
                    obs_entrega = st.text_area("Observação (opcional)", key=f"obs_entrega_{cautela['id']}")
                    if st.button("✅ Confirmar Entrega", key=f"confirmar_entrega_{cautela['id']}", use_container_width=True):
                        if entregar_cautela(cautela["id"], obs_entrega):
                            st.success("Material entregue! Cautela encerrada.")
                            st.rerun()
                        else:
                            st.error("Falha ao registrar a entrega — sem conexão com o banco.")

        ajuda("💡 Para imprimir uma cautela, use o botão \"IMPRIMIR / SALVAR ABA EM PDF\" no final da página.")

    st.divider()

    # Bloco 09: Alterações de Serviço (Antigo Bloco 06)
    st.markdown("### 09 - ALTERAÇÕES DE SERVIÇO")
    alteracoes_servico = campo_texto_com_voz("Descreva detalhadamente as alterações que venham a ocorrer durante o serviço", "alteracoes_servico_input", placeholder="Ex: Sem alterações relevantes...")
    st.divider()


    if encerrar_servico:
        texto_botao = "✅ ENCERRAR SERVIÇO E ENVIAR RELATÓRIO FINAL PARA A NUVEM"
    else:
        texto_botao = "💾 SALVAR PROGRESSO DO DIA (o serviço continua em andamento)"

    enviar = st.button(texto_botao, type="primary", use_container_width=True)

    # Todo este bloco só executa quando o botão é de fato clicado — antes ele
    # rodava a cada interação da tela (bug corrigido), o que tentava gravar
    # um registro novo no banco a cada tecla digitada.
    if enviar and not st.session_state["relatorio_enviado"]:
        if encerrar_servico and km_final < km_inicial:
            st.error("Erro: O KM Final não pode ser menor que o KM Inicial.")
        else:
            st.session_state["relatorio_enviado"] = True
            _id_existente_antes = st.session_state["relatorio_id_atual"]

            dados_para_salvar = {
                "status": "Finalizado" if encerrar_servico else "Em Andamento",
                "unidade": unidade_sel,
                "equipe": equipe_sel,
                "finalidade": finalidade_sel,
                "comandante": comandante_sel,
                "motorista": motorista_sel,
                "data_inicial": data_ini_sel,
                "data_final": data_fim_sel,
                "viatura_prefixo": viatura,
                "observacao_viatura": observacao_viatura,
                "km_inicial": km_inicial,
                "km_final": km_final if encerrar_servico else None,
                "capturas_animais": json.dumps(st.session_state["capturas_animais_list"], ensure_ascii=False),
                "apreensoes": json.dumps(st.session_state["apreensoes_list"], ensure_ascii=False),
                "visitas_prolepse": json.dumps(st.session_state["prolepse_list"], ensure_ascii=False),
                "pessoas_abordadas": sum(item["PESSOAS ABORDADAS"] for item in st.session_state["patrulhamento_terrestre_list"]),
                "veiculos_abordados": sum(item.get("VEÍCULOS ABORDADOS", 0) for item in st.session_state["patrulhamento_terrestre_list"]),
                "embarcacoes_abordadas": sum(item["EMBARCAÇÕES ABORDADAS"] for item in st.session_state["patrulhamento_fluvial_list"]),
                "bo_lavrados": 0,
                "autos_infracao": sum(len([v for v in ap.get("AUTO INFRAÇÃO", "").split("; ") if v.strip()]) for ap in st.session_state["apreensoes_list"]),
                "prolepse": prolepse if 'prolepse' in locals() else 0,
                "patrulhamento_terrestre": json.dumps(st.session_state["patrulhamento_terrestre_list"], ensure_ascii=False),
                "patrulhamento_fluvial": json.dumps(st.session_state["patrulhamento_fluvial_list"], ensure_ascii=False),
                "armamento_municao": json.dumps(tabela_armamento, ensure_ascii=False) if 'tabela_armamento' in locals() else "[]",
                "alteracoes_servico": alteracoes_servico,
                "ordens_servico_citadas": "; ".join(ordens_citadas_sel),
            }

            novo_id, erro = salvar_relatorio(dados_para_salvar, st.session_state["relatorio_id_atual"])

            if erro:
                st.session_state["relatorio_enviado"] = False
                st.error(f"Falha ao salvar no banco Neon: {erro}")
            else:
                st.session_state["relatorio_id_atual"] = novo_id
                registrar_auditoria(
                    novo_id, "servico",
                    "Edição" if _id_existente_antes else "Criação",
                    st.session_state.get("usuario_conectado", "")
                )
                if encerrar_servico:
                    # Serviço concluído: NÃO limpa mais aqui — abre o modal e deixa
                    # os dados na tela até o usuário escolher Imprimir e/ou Concluir.
                    # Isso corrige o bug de impressão em branco (a limpeza acontecia
                    # antes de a pessoa conseguir imprimir).
                    st.session_state["relatorio_aguardando_acao"] = novo_id
                else:
                    st.success(f"💾 Progresso do dia salvo na nuvem (Nº {novo_id:04d}). Serviço continua em andamento — pode fechar o sistema com segurança e retomar depois.")
                st.session_state["relatorio_enviado"] = False
                time.sleep(1)
                st.rerun()

    st.divider()
    _meses_pt = ["", "janeiro", "fevereiro", "março", "abril", "maio", "junho",
                 "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]
    _data_assinatura_txt = f"Miranda - MS, {data_fim_sel.day:02d} de {_meses_pt[data_fim_sel.month]} de {data_fim_sel.year}."
    st.markdown(f'<p style="text-align:right; margin:0;">{_data_assinatura_txt}</p>', unsafe_allow_html=True)
    _numero_rsd = st.session_state['relatorio_id_atual'] if st.session_state["relatorio_id_atual"] else proximo_numero
    st.markdown(f"""
        <div style="margin-top:20px; padding-top:10px; text-align:center;">
            <p style="margin-bottom:40px;">____________________________________________</p>
            <p style="margin:0; font-weight:bold; font-size:16px;">{comandante_sel}</p>
            <p style="margin:0;">Matrícula: {matricula_comandante if matricula_comandante else '____________________'}</p>
            <p style="margin:0; color:#888;">Comandante da Guarnição</p>
            <p style="margin-top:20px; color:#888; font-size:0.85em;">RSD Nº {_numero_rsd:04d}</p>
        </div>
    """, unsafe_allow_html=True)

    st.divider()
    # Este botão continua existindo só porque a impressão de CAUTELA (armamento/munição)
    # ainda depende dele — não passa pelo modal de finalização. Para o Relatório de
    # Serviço em si, use o botão "🖨️ Imprimir / Salvar em PDF" que aparece no modal
    # ao clicar em "Encerrar o serviço agora".
    ajuda("🖨️ Botão de impressão geral (usado principalmente para imprimir cautelas de armamento — o Relatório de Serviço tem seu próprio botão de impressão no modal ao finalizar).")
    components.html(
        """
        <button onclick="window.parent.print()" style="width:100%; padding:10px; background-color:#ff4b4b; color:white; border:none; border-radius:5px; cursor:pointer; font-weight:bold; font-size:15px;">
            IMPRIMIR / SALVAR ABA EM PDF
        </button>
        """,
        height=55,
    )

# ------------------------------------------
# VISÃO 2: PAINEL ESTRATÉGICO ADMINISTRATIVO (COMPLETO)
# ------------------------------------------
# ------------------------------------------
# VISÃO 3: RELATÓRIO DE FISCALIZAÇÃO AMBIENTAL
# ------------------------------------------
with aba_fiscalizacao:
    st.markdown("# 📋 RELATÓRIO DE FISCALIZAÇÃO AMBIENTAL")
    st.caption(f"Unidade: **{st.session_state['unidade_operacional']}**")

    if "rf_fotos_list" not in st.session_state:
        st.session_state["rf_fotos_list"] = []
    if "rf_id_atual" not in st.session_state:
        st.session_state["rf_id_atual"] = None
    if "rf_numero_atual" not in st.session_state:
        st.session_state["rf_numero_atual"] = None

    # --- Carrega um relatório encontrado na busca (antes dos widgets) ---
    if st.session_state.get("carregar_edicao_rf") is not None:
        _reg = st.session_state.pop("carregar_edicao_rf")
        st.session_state["rf_id_atual"] = _reg["id"]
        st.session_state["rf_numero_atual"] = _reg.get("numero")
        _interessado_carregado = _reg.get("interessado", "") or ""
        st.session_state["rf_interessado_multi"] = [v.strip() for v in _interessado_carregado.split(";") if v.strip()]
        st.session_state["rf_nome_autuado"] = _reg.get("nome_autuado", "")
        st.session_state["rf_cpf_cnpj"] = _reg.get("cpf_cnpj", "")
        st.session_state["rf_rg_ie"] = _reg.get("rg_ie", "")
        st.session_state["rf_endereco"] = _reg.get("endereco", "")
        st.session_state["rf_local"] = _reg.get("local_fiscalizacao", "")
        if _reg.get("data_fiscalizacao"):
            st.session_state["rf_data_fiscalizacao"] = _reg.get("data_fiscalizacao")
        st.session_state["rf_geo_formato"] = "Graus Decimais (DD)"
        st.session_state["rf_geo_lat_dd"] = float(_reg.get("latitude") or 0.0)
        st.session_state["rf_geo_lon_dd"] = float(_reg.get("longitude") or 0.0)
        st.session_state["rf_municipio"] = _reg.get("municipio", "Miranda")
        st.session_state["rf_telefone"] = _reg.get("telefone", "")
        st.session_state["rf_legislacao"] = _reg.get("legislacao", "")
        st.session_state["rf_auto_infracao_nr"] = _reg.get("auto_infracao_nr", "")
        st.session_state["rf_laudo_constatacao_nr"] = _reg.get("laudo_constatacao_nr", "")
        st.session_state["rf_termo_paralisacao_nr"] = _reg.get("termo_paralisacao_nr", "")
        st.session_state["rf_notificacao_nr"] = _reg.get("notificacao_nr", "")
        st.session_state["rf_folhas_complementares"] = _reg.get("folhas_complementares", 0) or 0
        st.session_state["rf_bo_cadg_nr"] = _reg.get("bo_cadg_nr", "")
        st.session_state["rf_fatos_historico"] = _reg.get("fatos_historico", "")
        st.session_state["rf_valor_multa_texto"] = _reg.get("valor_multa_texto", "")
        st.session_state["rf_providencias"] = _reg.get("providencias", "")
        st.session_state["rf_municipio_assinatura"] = _reg.get("municipio_assinatura", "Miranda (MS)")
        if _reg.get("data_assinatura"):
            st.session_state["rf_data_assinatura"] = _reg.get("data_assinatura")
        st.session_state["rf_relator"] = _reg.get("relator", "")
        st.session_state["rf_cargo_relator"] = _reg.get("cargo_relator", "Cmt. da GU Ambiental/Relator")
        try:
            st.session_state["rf_fotos_list"] = json.loads(_reg.get("fotos") or "[]")
        except Exception:
            st.session_state["rf_fotos_list"] = []

    # --- Busca dentro da própria aba (só nesta unidade) ---
    with st.expander("🔎 Buscar Relatório de Fiscalização (nesta unidade)"):
        with st.form("form_busca_rf"):
            termo_rf = st.text_input("Buscar por número ou nome do autuado", key="termo_busca_rf")
            buscar_rf_clicado = st.form_submit_button("🔎 Buscar")
        if buscar_rf_clicado and termo_rf:
            resultados_rf = buscar_relatorios_fiscalizacao(termo_rf, unidade=st.session_state["unidade_operacional"])
            if resultados_rf:
                for reg in resultados_rf:
                    col_r1, col_r2 = st.columns([5, 1])
                    with col_r1:
                        st.write(f"**{reg.get('numero','')}** — {reg.get('nome_autuado','')} ({reg.get('data_fiscalizacao','')}) — *{reg.get('status','Finalizado')}*")
                    if reg.get("status") == "Finalizado":
                        senha_desbloqueio_rf = st.text_input(
                            "Senha do administrador para editar", type="password",
                            key=f"senha_desbloqueio_rf_{reg['id']}", label_visibility="collapsed",
                            placeholder="Senha do administrador para editar este relatório"
                        )
                        if st.button("🔓 Desbloquear e Editar", key=f"desbloquear_rf_{reg['id']}", use_container_width=True):
                            if senha_desbloqueio_rf == USUARIOS_PERMITIDOS.get("admin"):
                                st.session_state["carregar_edicao_rf"] = reg
                                registrar_auditoria(reg['id'], "fiscalizacao", "Liberação de Edição (Admin)", st.session_state.get("usuario_conectado", ""))
                                st.rerun()
                            else:
                                st.error("Senha de administrador incorreta. Edição não liberada.")
                    else:
                        with col_r2:
                            if st.button("👁️ Abrir", key=f"abrir_rf_{reg['id']}", use_container_width=True):
                                st.session_state["carregar_edicao_rf"] = reg
                                st.rerun()
            else:
                st.info("Nenhum relatório encontrado.")

    st.divider()

    if st.session_state["rf_id_atual"]:
        col_novo1, col_novo2 = st.columns([4, 1])
        with col_novo1:
            st.info(f"✏️ Editando o relatório **Nº {st.session_state['rf_numero_atual']}**")
        with col_novo2:
            if st.button("➕ Novo Relatório", use_container_width=True, key="novo_rf_btn"):
                for k in [k for k in st.session_state.keys() if k.startswith("rf_")]:
                    del st.session_state[k]
                st.session_state["rf_fotos_list"] = []
                st.session_state["rf_id_atual"] = None
                st.session_state["rf_numero_atual"] = None
                st.rerun()

    lista_interessados_padrao = [
        "Instituto de Meio Ambiente de Mato Grosso do Sul (IMASUL)",
        "Ministério Público Estadual (MPE)",
        "CPAmb/1ºBPMA"
    ]
    if "rf_interessado_multi" not in st.session_state:
        st.session_state["rf_interessado_multi"] = lista_interessados_padrao.copy()

    interessados_sel_rf = st.multiselect(
        "Interessado(s)",
        options=lista_interessados_padrao,
        key="rf_interessado_multi"
    )
    interessado_rf = "; ".join(interessados_sel_rf)
    st.markdown("#### DO AUTUADO/FISCALIZADO")
    nome_autuado_rf = st.text_input("01 - Nome/Nome Empresarial", key="rf_nome_autuado")
    col_rf1, col_rf2 = st.columns(2)
    with col_rf1:
        cpf_cnpj_rf = st.text_input("02 - CPF/CNPJ", key="rf_cpf_cnpj")
    with col_rf2:
        rg_ie_rf = st.text_input("03 - RG/Insc. Estadual", key="rf_rg_ie")
    endereco_rf = st.text_area("04 - Endereço Completo", key="rf_endereco", height=70)
    st.markdown(
        f'<div class="print-only-text"><strong>Endereço Completo:</strong><br>{html.escape(endereco_rf or "—").replace(chr(10), "<br>")}</div>',
        unsafe_allow_html=True
    )
    st.markdown("#### DA INFRAÇÃO/FISCALIZAÇÃO")
    col_rf3, col_rf4 = st.columns(2)
    with col_rf3:
        local_rf = st.text_area("05 - Local", key="rf_local", height=70)
        st.markdown(
            f'<div class="print-only-text"><strong>Local:</strong><br>{html.escape(local_rf or "—").replace(chr(10), "<br>")}</div>',
            unsafe_allow_html=True
        )
    with col_rf4:
        data_fiscalizacao_rf = st.date_input("06 - Data", value=datetime.now(), key="rf_data_fiscalizacao")
    col_rf6, col_rf7 = st.columns(2)
    with col_rf6:
        municipio_rf = st.selectbox("08 - Município", ["Miranda", "Bodoquena", "Anastácio", "Aquidauana", "Corumbá", "Bonito", "Jardim", "Outros"], key="rf_municipio")
    with col_rf7:
        telefone_rf = st.text_input("09 - Telefone", key="rf_telefone")

    lat_rf, lon_rf, coordenadas_rf, area_geojson_rf = campo_geolocalizacao("rf_geo", permitir_area=True)

    st.markdown("#### LEGISLAÇÃO APLICÁVEL")
    legislacao_rf = campo_texto_com_voz("10 - Legislação (um artigo por linha)", "rf_legislacao", altura=100)

    st.markdown("#### FORMULÁRIOS IMASUL")
    col_rf8, col_rf9 = st.columns(2)
    with col_rf8:
        auto_infracao_nr_rf = st.text_input("Auto de Infração nº", key="rf_auto_infracao_nr")
    with col_rf9:
        laudo_constatacao_nr_rf = st.text_input("Laudo de Constatação nº", key="rf_laudo_constatacao_nr")
    col_rf10, col_rf11 = st.columns(2)
    with col_rf10:
        termo_paralisacao_nr_rf = st.text_input("Termo de Paralisação nº", key="rf_termo_paralisacao_nr")
    with col_rf11:
        notificacao_nr_rf = st.text_input("Notificação nº", key="rf_notificacao_nr")
    col_rf12, col_rf13 = st.columns(2)
    with col_rf12:
        folhas_complementares_rf = st.number_input("Folhas Complementares (quantidade)", min_value=0, step=1, key="rf_folhas_complementares")
    with col_rf13:
        bo_cadg_nr_rf = st.text_input("BO CADG nº", key="rf_bo_cadg_nr")

    st.markdown("#### 1. DOS FATOS — 1. HISTÓRICO")
    fatos_historico_rf = campo_texto_com_voz("Histórico", "rf_fatos_historico", altura=220)

    st.markdown("##### Fotos")
    novas_fotos = st.file_uploader("Adicionar fotos", type=["png", "jpg", "jpeg"], accept_multiple_files=True, key="rf_fotos_upload")
    if novas_fotos:
        for arquivo_foto in novas_fotos:
            try:
                dados_b64 = base64.b64encode(arquivo_foto.getvalue()).decode("utf-8")
                if not any(f.get("nome") == arquivo_foto.name and f.get("dados_base64") == dados_b64 for f in st.session_state["rf_fotos_list"]):
                    st.session_state["rf_fotos_list"].append({"nome": arquivo_foto.name, "legenda": "", "dados_base64": dados_b64})
            except Exception:
                pass

    for i_foto, foto in enumerate(st.session_state["rf_fotos_list"]):
        col_f1, col_f2, col_f3 = st.columns([2, 5, 1])
        with col_f1:
            try:
                st.image(base64.b64decode(foto["dados_base64"]), width=100)
            except Exception:
                st.caption("(imagem)")
        with col_f2:
            foto["legenda"] = st.text_input(f"Legenda da foto {i_foto+1}", value=foto.get("legenda", ""), key=f"rf_legenda_foto_{i_foto}", label_visibility="collapsed", placeholder=f"Ex: Foto {i_foto+1}: descrição...")
        with col_f3:
            if st.button("🗑️", key=f"rf_del_foto_{i_foto}"):
                st.session_state["rf_fotos_list"].pop(i_foto)
                st.rerun()

    st.markdown("#### 2. DO VALOR DA MULTA")
    valor_multa_texto_rf = campo_texto_com_voz("Texto sobre a multa", "rf_valor_multa_texto", altura=100)

    st.markdown("#### 3. DAS PROVIDÊNCIAS ADMINISTRATIVAS")
    providencias_rf = campo_texto_com_voz("Providências (uma por linha)", "rf_providencias", altura=150)

    st.markdown("#### Assinatura")
    col_rf14, col_rf15 = st.columns(2)
    with col_rf14:
        municipio_assinatura_rf = st.text_input("Município (assinatura)", value=st.session_state.get("rf_municipio_assinatura", "Miranda (MS)"), key="rf_municipio_assinatura")
    with col_rf15:
        data_assinatura_rf = st.date_input("Data da Assinatura", value=datetime.now(), key="rf_data_assinatura")
    relator_rf = st.selectbox("Relator (Comandante)", EFETIVO[st.session_state["unidade_operacional"]], key="rf_relator")
    matricula_relator_rf = MATRICULAS.get(relator_rf, "")
    st.caption(f"Matrícula: {matricula_relator_rf or '— não cadastrada —'}")
    cargo_relator_rf = st.text_input("Cargo/Função", value=st.session_state.get("rf_cargo_relator", "Cmt. da GU Ambiental/Relator"), key="rf_cargo_relator")

    @st.dialog("Relatório de Fiscalização Salvo")
    def modal_rf_finalizado(id_relatorio, numero_relatorio):
        st.success(f"✅ Relatório {numero_relatorio} finalizado e gravado com sucesso na nuvem do Pelotão!")
        st.caption("Os dados continuam na tela até você clicar em Concluir — pode imprimir quantas vezes precisar antes disso.")

        if st.session_state.get("_disparar_impressao_rf"):
            components.html("<script>window.parent.print();</script>", height=0)
            st.session_state["_disparar_impressao_rf"] = False

        col_rfm1, col_rfm2 = st.columns(2)
        with col_rfm1:
            if st.button("🖨️ Imprimir / Salvar em PDF", use_container_width=True, key="rf_modal_imprimir"):
                registrar_auditoria(id_relatorio, "fiscalizacao", "Impressão", st.session_state.get("usuario_conectado", ""))
                st.session_state["_disparar_impressao_rf"] = True
                st.rerun()
        with col_rfm2:
            if st.button("✅ Concluir e Iniciar Novo Relatório", type="primary", use_container_width=True, key="rf_modal_concluir"):
                for _k in list(st.session_state.keys()):
                    if _k.startswith("rf_") and _k not in ("rf_geo_formato",):
                        del st.session_state[_k]
                st.session_state["relatorio_rf_aguardando_acao"] = None
                st.rerun()

    if st.session_state.get("relatorio_rf_aguardando_acao"):
        _dados_modal_rf = st.session_state["relatorio_rf_aguardando_acao"]
        modal_rf_finalizado(_dados_modal_rf["id"], _dados_modal_rf["numero"])

    col_salvar_rf1, col_salvar_rf2 = st.columns(2)
    with col_salvar_rf1:
        salvar_rascunho_rf = st.button("💾 Salvar Rascunho (continuar depois)", use_container_width=True)
    with col_salvar_rf2:
        finalizar_rf = st.button("✅ Finalizar Relatório de Fiscalização", type="primary", use_container_width=True)

    if salvar_rascunho_rf or finalizar_rf:
        dados_rf = {
            "interessado": interessado_rf, "nome_autuado": nome_autuado_rf, "cpf_cnpj": cpf_cnpj_rf,
            "rg_ie": rg_ie_rf, "endereco": endereco_rf, "local_fiscalizacao": local_rf,
            "data_fiscalizacao": data_fiscalizacao_rf, "coordenadas": coordenadas_rf, "municipio": municipio_rf,
            "latitude": lat_rf, "longitude": lon_rf, "area_poligono": area_geojson_rf,
            "telefone": telefone_rf, "legislacao": legislacao_rf, "auto_infracao_nr": auto_infracao_nr_rf,
            "laudo_constatacao_nr": laudo_constatacao_nr_rf, "termo_paralisacao_nr": termo_paralisacao_nr_rf,
            "notificacao_nr": notificacao_nr_rf, "folhas_complementares": folhas_complementares_rf,
            "bo_cadg_nr": bo_cadg_nr_rf, "fatos_historico": fatos_historico_rf,
            "fotos": json.dumps(st.session_state["rf_fotos_list"], ensure_ascii=False),
            "valor_multa_texto": valor_multa_texto_rf, "providencias": providencias_rf,
            "municipio_assinatura": municipio_assinatura_rf, "data_assinatura": data_assinatura_rf,
            "relator": relator_rf, "matricula_relator": matricula_relator_rf, "cargo_relator": cargo_relator_rf,
            "status": "Finalizado" if finalizar_rf else "Em Andamento"
        }
        if not st.session_state["rf_id_atual"]:
            ano_rf = data_fiscalizacao_rf.year
            sequencial_rf = proximo_numero_fiscalizacao(st.session_state["unidade_operacional"], ano_rf)
            designacao_rf = DESIGNACAO_UNIDADE_FISCALIZACAO.get(st.session_state["unidade_operacional"], "")
            numero_rf = formatar_numero_fiscalizacao(sequencial_rf, designacao_rf, ano_rf)
            dados_rf["numero"] = numero_rf
            dados_rf["sequencial"] = sequencial_rf
            dados_rf["ano"] = ano_rf
            dados_rf["unidade"] = st.session_state["unidade_operacional"]

        novo_id_rf, erro_rf = salvar_relatorio_fiscalizacao(
            dados_rf, id_existente=st.session_state["rf_id_atual"], usuario_atual=st.session_state.get("usuario_conectado", "")
        )
        if erro_rf:
            st.error(f"Falha ao salvar: {erro_rf}")
        else:
            st.session_state["rf_id_atual"] = novo_id_rf
            if not st.session_state.get("rf_numero_atual"):
                st.session_state["rf_numero_atual"] = dados_rf.get("numero")

            if finalizar_rf:
                st.session_state["relatorio_rf_aguardando_acao"] = {"id": novo_id_rf, "numero": st.session_state["rf_numero_atual"]}
                st.rerun()
            else:
                st.success(f"💾 Rascunho salvo (Nº {st.session_state['rf_numero_atual']}). Pode fechar o sistema com segurança — os dados já estão na nuvem, use a busca acima para retomar depois.")

            dados_rf["numero"] = st.session_state["rf_numero_atual"]
            buffer_docx = gerar_docx_fiscalizacao(dados_rf)
            st.download_button(
                "⬇️ Baixar Documento (.docx)", buffer_docx,
                file_name=f"Relatorio_Fiscalizacao_{st.session_state['rf_numero_atual'].replace('/', '-')}.docx",
                use_container_width=True
            )

# ------------------------------------------
# VISÃO 2: PAINEL ESTRATÉGICO (ADMIN)
# ------------------------------------------
with aba_adm:
    st.markdown("# PAINEL GERENCIAL DA ADMINISTRAÇÃO")
    senha = st.text_input("Insira a senha administrativa", type="password", key="input_senha")
    
    if senha == "adm123":
        st.success("Acesso liberado!")
        conn = init_connection()
        
        if conn:
            try:
                cur = conn.cursor()
                cur.execute("SELECT * FROM relatorios_servico ORDER BY id ASC;")
                registros = cur.fetchall()
                cur.close()
                conn.close()
                
                if registros:
                    df = pd.DataFrame(registros)
                    df.columns = [c.lower() for c in df.columns]
                    
                    # Conversões e tratamentos numéricos
                    df['pessoas_abordadas'] = pd.to_numeric(df['pessoas_abordadas'], errors='coerce').fillna(0)
                    df['autos_infracao'] = pd.to_numeric(df['autos_infracao'], errors='coerce').fillna(0)
                    df['km_rodado'] = 0.0
                    if 'status' in df.columns:
                        _mask_final = df['status'] == 'Finalizado'
                        df.loc[_mask_final, 'km_rodado'] = pd.to_numeric(df.loc[_mask_final, 'km_final'], errors='coerce').fillna(0) - pd.to_numeric(df.loc[_mask_final, 'km_inicial'], errors='coerce').fillna(0)
                    df['km_rodado'] = df['km_rodado'].clip(lower=0)
                    
                    if 'quantidade_apreendida' in df.columns:
                        df['quantidade_apreendida'] = pd.to_numeric(df['quantidade_apreendida'], errors='coerce').fillna(0)
                    else:
                        df['quantidade_apreendida'] = 0
                    if 'valor_multa' in df.columns:
                        df['valor_multa'] = pd.to_numeric(df['valor_multa'], errors='coerce').fillna(0)
                    else:
                        df['valor_multa'] = 0

                    for col_json in ['patrulhamento_terrestre', 'patrulhamento_fluvial', 'capturas_animais', 'apreensoes']:
                        if col_json not in df.columns:
                            df[col_json] = "[]"

                    # Novos campos de documentos da apreensão (Nº CADG, Auto de Infração, etc.)
                    for col_doc in ['cadg', 'nr_auto_infracao', 'nr_termo_constatacao', 'nr_termo_apreensao']:
                        if col_doc not in df.columns:
                            df[col_doc] = ""
                        else:
                            df[col_doc] = df[col_doc].fillna("")
                    
                    # Tratamento de datas para o filtro de tempo
                    if 'data_criacao' in df.columns:
                        df['data_filtro'] = pd.to_datetime(df['data_criacao']).dt.date
                    elif 'data_inicial' in df.columns:
                        df['data_filtro'] = pd.to_datetime(df['data_inicial']).dt.date
                    else:
                        df['data_filtro'] = datetime.now().date()
                    
                    # Renomeia o ID para Sequência Numérica Oficial
                    df = df.rename(columns={"id": "Nº Sequencial do Relatório"})

                    # Cópia sem filtro de período — usada só pela Busca Avançada por ano abaixo
                    df_completo_todos_anos = df.copy()

                    # Filtros de tempo (Semana, Mês, Ano)
                    st.markdown("### 📅 Filtro Temporal de Produção")
                    periodo = st.radio("Selecione o período de análise:", ["Total Histórico", "Últimos 7 Dias (Semana)", "Último Mês", "Ano Atual"], horizontal=True)
                    
                    hoje = datetime.now().date()
                    if periodo == "Últimos 7 Dias (Semana)":
                        df = df[df['data_filtro'] >= (hoje - timedelta(days=7))]
                    elif periodo == "Último Mês":
                        df = df[df['data_filtro'] >= (hoje - timedelta(days=30))]
                    elif periodo == "Ano Atual":
                        df = df[pd.to_datetime(df['data_filtro']).dt.year == datetime.now().year]

                    def explodir_lista_json(df_base, coluna_json, mapa_colunas=None):
                        """Transforma a coluna JSON (lista de dicts, uma por atividade/captura)
                        de cada relatório em um DataFrame 'longo', uma linha por item, associado
                        ao Nº do relatório de origem."""
                        linhas = []
                        for _, row in df_base.iterrows():
                            try:
                                itens = json.loads(row.get(coluna_json) or "[]")
                            except Exception:
                                itens = []
                            for item in itens:
                                item_com_ref = dict(item)
                                if mapa_colunas:
                                    for origem, destino in mapa_colunas.items():
                                        if origem in item_com_ref:
                                            item_com_ref[destino] = item_com_ref.pop(origem)
                                item_com_ref["Nº Relatório"] = row.get("Nº Sequencial do Relatório")
                                linhas.append(item_com_ref)
                        return pd.DataFrame(linhas) if linhas else pd.DataFrame()

                    df_graf_terr = explodir_lista_json(df, "patrulhamento_terrestre", {"PESSOAS ABORDADAS": "Pessoas Abordadas", "APREENSÕES": "Apreensões"})
                    df_graf_fluv = explodir_lista_json(df, "patrulhamento_fluvial", {"PESCADORES ABORDADOS": "Pescadores Abordados", "APREENSÕES": "Apreensões", "EMBARCAÇÕES ABORDADAS": "Embarcações Abordadas"})
                    df_graf_animais = explodir_lista_json(df, "capturas_animais")
                    if not df_graf_animais.empty and "QUANTIDADE" in df_graf_animais.columns:
                        total_animais_capturados = int(pd.to_numeric(df_graf_animais["QUANTIDADE"], errors='coerce').fillna(0).sum())
                    else:
                        total_animais_capturados = 0

                    df_graf_apreensoes = explodir_lista_json(df, "apreensoes")
                    if not df_graf_apreensoes.empty:
                        df_graf_apreensoes = df_graf_apreensoes[df_graf_apreensoes["INFRAÇÃO/CRIME"] != "Não se aplica"]
                    total_apreensoes = len(df_graf_apreensoes) if not df_graf_apreensoes.empty else 0
                    total_multas = pd.to_numeric(df_graf_apreensoes["VALOR MULTA"], errors='coerce').fillna(0).sum() if not df_graf_apreensoes.empty and "VALOR MULTA" in df_graf_apreensoes.columns else 0.0

                    # --- BUSCA AVANÇADA: as duas unidades juntas, com seletor de ano ---
                    st.divider()
                    with st.expander("🔎 Busca Avançada (todas as unidades, por ano de referência)"):
                        anos_com_dados = []
                        if not df_completo_todos_anos.empty:
                            anos_com_dados = pd.to_datetime(df_completo_todos_anos['data_filtro']).dt.year.dropna().unique().tolist()
                        ano_atual = datetime.now().year
                        anos_disponiveis = sorted(set([int(a) for a in anos_com_dados] + [ano_atual, ano_atual + 1]), reverse=True)

                        with st.form("form_busca_avancada"):
                            col_busca1, col_busca2 = st.columns([1, 3])
                            with col_busca1:
                                ano_busca = st.selectbox(
                                    "Ano de referência", anos_disponiveis, key="ano_busca_avancada",
                                    index=anos_disponiveis.index(ano_atual) if ano_atual in anos_disponiveis else 0
                                )
                            with col_busca2:
                                termo_busca_avancada = st.text_input(
                                    "Buscar (ex: nome de uma espécie de animal, Nº do relatório, comandante, tipo de material...)",
                                    key="termo_busca_avancada"
                                )
                            buscar_clicado = st.form_submit_button("🔎 Buscar")

                        if buscar_clicado:
                            st.session_state["ultima_busca_avancada"] = {"ano": ano_busca, "termo": termo_busca_avancada}

                        busca_ativa = st.session_state.get("ultima_busca_avancada")

                        if busca_ativa and busca_ativa["termo"]:
                            ano_busca_ativa = busca_ativa["ano"]
                            termo_busca_avancada = busca_ativa["termo"]
                            df_ano_busca = df_completo_todos_anos[pd.to_datetime(df_completo_todos_anos['data_filtro']).dt.year == ano_busca_ativa]
                            termo_lower = termo_busca_avancada.lower()

                            df_animais_ano = explodir_lista_json(df_ano_busca, "capturas_animais")
                            if not df_animais_ano.empty and "ESPÉCIE" in df_animais_ano.columns:
                                resultado_animais = df_animais_ano[df_animais_ano["ESPÉCIE"].astype(str).str.lower().str.contains(termo_lower, na=False)]
                            else:
                                resultado_animais = pd.DataFrame()

                            df_apreensoes_ano = explodir_lista_json(df_ano_busca, "apreensoes")
                            if not df_apreensoes_ano.empty:
                                mask_apreensao = pd.Series(False, index=df_apreensoes_ano.index)
                                for col in ["INFRAÇÃO/CRIME", "TIPO MATERIAL", "DESCRIÇÃO"]:
                                    if col in df_apreensoes_ano.columns:
                                        mask_apreensao |= df_apreensoes_ano[col].astype(str).str.lower().str.contains(termo_lower, na=False)
                                resultado_apreensoes = df_apreensoes_ano[mask_apreensao]
                            else:
                                resultado_apreensoes = pd.DataFrame()

                            mask_relatorio = pd.Series(False, index=df_ano_busca.index)
                            for col in ["Nº Sequencial do Relatório", "comandante", "unidade"]:
                                if col in df_ano_busca.columns:
                                    mask_relatorio |= df_ano_busca[col].astype(str).str.lower().str.contains(termo_lower, na=False)
                            resultado_relatorios = df_ano_busca[mask_relatorio]

                            achou_algo = False
                            if not resultado_animais.empty:
                                achou_algo = True
                                qtd_total_animais = pd.to_numeric(resultado_animais["QUANTIDADE"], errors='coerce').fillna(0).sum() if "QUANTIDADE" in resultado_animais.columns else len(resultado_animais)
                                st.success(f"🐾 **{qtd_total_animais:.0f}** capturado(s) envolvendo \"{termo_busca_avancada}\" em {ano_busca_ativa} ({len(resultado_animais)} ocorrência(s))")
                                st.dataframe(resultado_animais, use_container_width=True)

                            if not resultado_apreensoes.empty:
                                achou_algo = True
                                st.info(f"📦 **{len(resultado_apreensoes)}** apreensão(ões) relacionada(s) a \"{termo_busca_avancada}\" em {ano_busca_ativa}")
                                st.dataframe(resultado_apreensoes, use_container_width=True)

                            if not resultado_relatorios.empty:
                                achou_algo = True
                                st.info(f"📋 **{len(resultado_relatorios)}** relatório(s) encontrados por Nº/comandante/unidade")
                                cols_rel = [c for c in ["Nº Sequencial do Relatório", "unidade", "comandante", "status", "data_filtro"] if c in resultado_relatorios.columns]
                                st.dataframe(resultado_relatorios[cols_rel], use_container_width=True)

                            if not achou_algo:
                                st.warning("Nenhum resultado encontrado para esse termo nesse ano.")
                        else:
                            st.caption("Escolha o ano, digite um termo de busca e clique em \"🔎 Buscar\".")

                    # --- GERENCIAMENTO DE RELATÓRIOS DE FISCALIZAÇÃO AMBIENTAL (as duas unidades) ---
                    st.divider()
                    with st.expander("📋 Relatórios de Fiscalização Ambiental (buscar, visualizar, editar, excluir)"):
                        if st.session_state.get("carregar_edicao_rf_admin") is not None:
                            _reg_adm = st.session_state.pop("carregar_edicao_rf_admin")
                            for _campo, _chave_ss in [
                                ("interessado", "rf_adm_interessado"), ("nome_autuado", "rf_adm_nome_autuado"),
                                ("cpf_cnpj", "rf_adm_cpf_cnpj"), ("rg_ie", "rf_adm_rg_ie"), ("endereco", "rf_adm_endereco"),
                                ("local_fiscalizacao", "rf_adm_local"), ("coordenadas", "rf_adm_coordenadas"),
                                ("municipio", "rf_adm_municipio"), ("telefone", "rf_adm_telefone"),
                                ("legislacao", "rf_adm_legislacao"), ("auto_infracao_nr", "rf_adm_auto_infracao_nr"),
                                ("laudo_constatacao_nr", "rf_adm_laudo_constatacao_nr"), ("termo_paralisacao_nr", "rf_adm_termo_paralisacao_nr"),
                                ("notificacao_nr", "rf_adm_notificacao_nr"), ("bo_cadg_nr", "rf_adm_bo_cadg_nr"),
                                ("fatos_historico", "rf_adm_fatos_historico"), ("valor_multa_texto", "rf_adm_valor_multa_texto"),
                                ("providencias", "rf_adm_providencias"), ("cargo_relator", "rf_adm_cargo_relator")
                            ]:
                                st.session_state[_chave_ss] = _reg_adm.get(_campo, "") or ""
                            st.session_state["rf_adm_folhas_complementares"] = _reg_adm.get("folhas_complementares", 0) or 0
                            st.session_state["rf_adm_id_atual"] = _reg_adm["id"]
                            st.session_state["rf_adm_numero_atual"] = _reg_adm.get("numero", "")
                            st.session_state["rf_adm_unidade_atual"] = _reg_adm.get("unidade", "")

                        with st.form("form_busca_rf_admin"):
                            termo_rf_adm = st.text_input("Buscar por número ou nome do autuado (todas as unidades)", key="termo_busca_rf_admin")
                            buscar_rf_admin_clicado = st.form_submit_button("🔎 Buscar")

                        if buscar_rf_admin_clicado and termo_rf_adm:
                            st.session_state["ultima_busca_rf_admin"] = termo_rf_adm

                        termo_rf_admin_ativo = st.session_state.get("ultima_busca_rf_admin")
                        if termo_rf_admin_ativo:
                            resultados_rf_admin = buscar_relatorios_fiscalizacao(termo_rf_admin_ativo)
                            if not resultados_rf_admin:
                                st.warning("Nenhum relatório encontrado.")
                            for reg_rf in resultados_rf_admin:
                                with st.container(border=True):
                                    st.markdown(f"**Nº {reg_rf.get('numero','')}** — {reg_rf.get('nome_autuado','')} | {reg_rf.get('unidade','')}")
                                    st.caption(f"Criado por: {reg_rf.get('criado_por','') or '—'}" + (f" · Editado por: {reg_rf.get('editado_por')} em {reg_rf.get('data_edicao')}" if reg_rf.get('editado_por') else ""))
                                    col_rfa1, col_rfa2, col_rfa3 = st.columns(3)
                                    with col_rfa1:
                                        if st.button("👁️ Visualizar/Editar", key=f"ver_rf_adm_{reg_rf['id']}", use_container_width=True):
                                            st.session_state["carregar_edicao_rf_admin"] = reg_rf
                                            st.rerun()
                                    with col_rfa2:
                                        try:
                                            _buffer_rf_view = gerar_docx_fiscalizacao(dict(reg_rf))
                                            st.download_button(
                                                "⬇️ Baixar .docx", _buffer_rf_view,
                                                file_name=f"Relatorio_Fiscalizacao_{reg_rf.get('numero','').replace('/', '-')}.docx",
                                                key=f"baixar_rf_adm_{reg_rf['id']}", use_container_width=True
                                            )
                                        except Exception:
                                            st.caption("Não foi possível gerar o .docx")
                                    with col_rfa3:
                                        with st.popover("🗑️ Excluir", use_container_width=True):
                                            st.warning("Essa ação é definitiva e não pode ser desfeita.")
                                            if st.button("Confirmar exclusão", key=f"confirmar_exclusao_rf_{reg_rf['id']}", use_container_width=True):
                                                if excluir_relatorio_fiscalizacao(reg_rf["id"]):
                                                    st.success("Relatório excluído.")
                                                    st.rerun()
                                                else:
                                                    st.error("Falha ao excluir — sem conexão com o banco.")

                        if st.session_state.get("rf_adm_id_atual"):
                            st.divider()
                            st.markdown(f"#### ✏️ Editando Relatório Nº {st.session_state.get('rf_adm_numero_atual','')}")
                            rf_adm_interessado = st.text_input("Interessado", key="rf_adm_interessado")
                            rf_adm_nome_autuado = st.text_input("01 - Nome/Nome Empresarial", key="rf_adm_nome_autuado")
                            col_rfa4, col_rfa5 = st.columns(2)
                            with col_rfa4:
                                rf_adm_cpf_cnpj = st.text_input("02 - CPF/CNPJ", key="rf_adm_cpf_cnpj")
                            with col_rfa5:
                                rf_adm_rg_ie = st.text_input("03 - RG/Insc. Estadual", key="rf_adm_rg_ie")
                            rf_adm_endereco = st.text_area("04 - Endereço Completo", key="rf_adm_endereco")
                            rf_adm_local = st.text_area("05 - Local", key="rf_adm_local")
                            rf_adm_coordenadas = st.text_input("07 - Coord. Geográfica", key="rf_adm_coordenadas")
                            rf_adm_municipio = st.text_input("08 - Município", key="rf_adm_municipio")
                            rf_adm_telefone = st.text_input("09 - Telefone", key="rf_adm_telefone")
                            rf_adm_legislacao = st.text_area("10 - Legislação", key="rf_adm_legislacao")
                            col_rfa6, col_rfa7 = st.columns(2)
                            with col_rfa6:
                                rf_adm_auto_infracao_nr = st.text_input("Auto de Infração nº", key="rf_adm_auto_infracao_nr")
                            with col_rfa7:
                                rf_adm_laudo_constatacao_nr = st.text_input("Laudo de Constatação nº", key="rf_adm_laudo_constatacao_nr")
                            col_rfa8, col_rfa9 = st.columns(2)
                            with col_rfa8:
                                rf_adm_termo_paralisacao_nr = st.text_input("Termo de Paralisação nº", key="rf_adm_termo_paralisacao_nr")
                            with col_rfa9:
                                rf_adm_notificacao_nr = st.text_input("Notificação nº", key="rf_adm_notificacao_nr")
                            col_rfa10, col_rfa11 = st.columns(2)
                            with col_rfa10:
                                rf_adm_folhas_complementares = st.number_input("Folhas Complementares", min_value=0, step=1, key="rf_adm_folhas_complementares")
                            with col_rfa11:
                                rf_adm_bo_cadg_nr = st.text_input("BO CADG nº", key="rf_adm_bo_cadg_nr")
                            rf_adm_fatos_historico = st.text_area("1. Histórico", key="rf_adm_fatos_historico", height=200)
                            rf_adm_valor_multa_texto = st.text_area("2. Valor da Multa", key="rf_adm_valor_multa_texto")
                            rf_adm_providencias = st.text_area("3. Providências Administrativas", key="rf_adm_providencias")
                            rf_adm_cargo_relator = st.text_input("Cargo/Função do Relator", key="rf_adm_cargo_relator")

                            if st.button("💾 Salvar Alterações", type="primary", key="salvar_rf_admin", use_container_width=True):
                                dados_rf_adm = {
                                    "interessado": rf_adm_interessado, "nome_autuado": rf_adm_nome_autuado,
                                    "cpf_cnpj": rf_adm_cpf_cnpj, "rg_ie": rf_adm_rg_ie, "endereco": rf_adm_endereco,
                                    "local_fiscalizacao": rf_adm_local, "coordenadas": rf_adm_coordenadas,
                                    "municipio": rf_adm_municipio, "telefone": rf_adm_telefone,
                                    "legislacao": rf_adm_legislacao, "auto_infracao_nr": rf_adm_auto_infracao_nr,
                                    "laudo_constatacao_nr": rf_adm_laudo_constatacao_nr, "termo_paralisacao_nr": rf_adm_termo_paralisacao_nr,
                                    "notificacao_nr": rf_adm_notificacao_nr, "folhas_complementares": rf_adm_folhas_complementares,
                                    "bo_cadg_nr": rf_adm_bo_cadg_nr, "fatos_historico": rf_adm_fatos_historico,
                                    "valor_multa_texto": rf_adm_valor_multa_texto, "providencias": rf_adm_providencias,
                                    "cargo_relator": rf_adm_cargo_relator
                                }
                                _id_rf_adm, _erro_rf_adm = salvar_relatorio_fiscalizacao(
                                    dados_rf_adm, id_existente=st.session_state["rf_adm_id_atual"],
                                    usuario_atual=st.session_state.get("usuario_conectado", "")
                                )
                                if _erro_rf_adm:
                                    st.error(f"Falha ao salvar: {_erro_rf_adm}")
                                else:
                                    st.success(f"✅ Alterações salvas — registrado como editado por {st.session_state.get('usuario_conectado','')}.")
                                    st.session_state["rf_adm_id_atual"] = None
                                    st.rerun()

                    # --- EXPORTAÇÃO EM LOTE: todos os relatórios de um período, num clique ---
                    st.divider()
                    with st.expander("📦 Exportação em Lote (Excel / Word consolidado por período)"):
                        col_exp1, col_exp2 = st.columns(2)
                        nomes_meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho",
                                       "Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
                        with col_exp1:
                            mes_exportar = st.selectbox(
                                "Mês", list(range(1, 13)), format_func=lambda m: nomes_meses[m - 1],
                                index=datetime.now().month - 1, key="exp_mes"
                            )
                        with col_exp2:
                            ano_exportar = st.selectbox(
                                "Ano", list(range(datetime.now().year - 3, datetime.now().year + 1))[::-1],
                                key="exp_ano"
                            )

                        data_ini_exp = datetime(ano_exportar, mes_exportar, 1).date()
                        if mes_exportar == 12:
                            data_fim_exp = datetime(ano_exportar, 12, 31).date()
                        else:
                            data_fim_exp = (datetime(ano_exportar, mes_exportar + 1, 1) - timedelta(days=1)).date()

                        df_periodo_servico = df_completo_todos_anos[
                            (df_completo_todos_anos['data_filtro'] >= data_ini_exp) &
                            (df_completo_todos_anos['data_filtro'] <= data_fim_exp)
                        ]

                        registros_rf_periodo = []
                        conn_exp = init_connection()
                        if conn_exp:
                            try:
                                cur_exp = conn_exp.cursor()
                                cur_exp.execute(
                                    "SELECT * FROM relatorios_fiscalizacao WHERE data_fiscalizacao >= %s AND data_fiscalizacao <= %s ORDER BY id ASC;",
                                    (data_ini_exp, data_fim_exp)
                                )
                                registros_rf_periodo = cur_exp.fetchall()
                                cur_exp.close()
                                conn_exp.close()
                            except Exception as e_exp:
                                st.warning(f"Não foi possível carregar os relatórios de fiscalização do período: {e_exp}")

                        st.caption(
                            f"Período: {data_ini_exp.strftime('%d/%m/%Y')} a {data_fim_exp.strftime('%d/%m/%Y')} — "
                            f"{len(df_periodo_servico)} relatório(s) de Serviço, {len(registros_rf_periodo)} de Fiscalização."
                        )

                        col_botao1, col_botao2 = st.columns(2)

                        with col_botao1:
                            if st.button("📊 Gerar Excel do Período", use_container_width=True, key="btn_gerar_excel"):
                                try:
                                    buffer_excel = io.BytesIO()
                                    with pd.ExcelWriter(buffer_excel, engine="openpyxl") as writer:
                                        colunas_servico_export = [c for c in [
                                            "Nº Sequencial do Relatório", "status", "data_filtro", "unidade", "finalidade",
                                            "comandante", "motorista", "viatura_prefixo", "km_rodado",
                                            "pessoas_abordadas", "veiculos_abordados", "embarcacoes_abordadas", "autos_infracao"
                                        ] if c in df_periodo_servico.columns]
                                        df_periodo_servico[colunas_servico_export].to_excel(writer, sheet_name="Relatórios de Serviço", index=False)

                                        if registros_rf_periodo:
                                            df_rf_export = pd.DataFrame(registros_rf_periodo)
                                            df_rf_export.columns = [c.lower() for c in df_rf_export.columns]
                                            colunas_rf_export = [c for c in [
                                                "id", "sequencial", "unidade", "data_fiscalizacao", "interessado", "nome_autuado",
                                                "municipio", "coordenadas", "auto_infracao_nr", "valor_multa_texto"
                                            ] if c in df_rf_export.columns]
                                            df_rf_export[colunas_rf_export].to_excel(writer, sheet_name="Relatórios de Fiscalização", index=False)
                                        else:
                                            pd.DataFrame().to_excel(writer, sheet_name="Relatórios de Fiscalização", index=False)
                                    buffer_excel.seek(0)
                                    st.download_button(
                                        "⬇️ Baixar Excel", data=buffer_excel,
                                        file_name=f"relatorios_{ano_exportar}_{mes_exportar:02d}.xlsx",
                                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                                        use_container_width=True, key="download_excel_lote"
                                    )
                                except Exception as e_excel:
                                    st.error(f"Falha ao gerar Excel: {e_excel}. Verifique se `openpyxl` está no requirements.txt.")

                        with col_botao2:
                            if st.button("📄 Gerar Word Consolidado (Fiscalização)", use_container_width=True, key="btn_gerar_word"):
                                if not registros_rf_periodo:
                                    st.info("Não há relatórios de fiscalização nesse período para consolidar.")
                                else:
                                    try:
                                        from docxcompose.composer import Composer
                                        doc_mestre = Document(gerar_docx_fiscalizacao(dict(registros_rf_periodo[0])))
                                        composer = Composer(doc_mestre)
                                        for reg_rf in registros_rf_periodo[1:]:
                                            sub_doc = Document(gerar_docx_fiscalizacao(dict(reg_rf)))
                                            composer.append(sub_doc)
                                        buffer_word = io.BytesIO()
                                        composer.save(buffer_word)
                                        buffer_word.seek(0)
                                        st.download_button(
                                            "⬇️ Baixar Word Consolidado", data=buffer_word,
                                            file_name=f"fiscalizacao_consolidado_{ano_exportar}_{mes_exportar:02d}.docx",
                                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                            use_container_width=True, key="download_word_lote"
                                        )
                                    except ImportError:
                                        st.error("Biblioteca `docxcompose` não instalada — adicione `docxcompose` ao requirements.txt e reimplante o app.")
                                    except Exception as e_word:
                                        st.error(f"Falha ao gerar o Word consolidado: {e_word}")
                        st.caption("O Word consolidado hoje reúne apenas os Relatórios de Fiscalização (é o único tipo com gerador de .docx pronto). O de Serviço sai completo na aba Excel.")

                    # Abas do Dashboard Administrativo
                    tab_geral, tab_unidades, tab_equipes = st.tabs(["📊 Produção Geral", "🏢 Por Unidade", "🪖 Por Equipes (Efetivo)"])
                    
                    with tab_geral:
                        m1, m2, m3, m4, m5 = st.columns(5)
                        m1.metric("Total de Abordagens", int(df['pessoas_abordadas'].sum()))
                        m2.metric("Total de KM Rodado", f"{int(df['km_rodado'].sum())} km")
                        m3.metric("Animais Computados", total_animais_capturados)
                        m4.metric("Apreensões", total_apreensoes)
                        m5.metric("Total em Multas", f"R$ {total_multas:,.2f}")

                        # --- Comparação com o mês anterior ---
                        st.divider()
                        st.markdown("### 📈 Comparação com o Mês Anterior")

                        hoje_cmp = datetime.now().date()
                        primeiro_dia_mes_atual = hoje_cmp.replace(day=1)
                        ultimo_dia_mes_anterior = primeiro_dia_mes_atual - timedelta(days=1)
                        primeiro_dia_mes_anterior = ultimo_dia_mes_anterior.replace(day=1)

                        def metricas_do_periodo(data_ini, data_fim):
                            df_p = df_completo_todos_anos[
                                (df_completo_todos_anos['data_filtro'] >= data_ini) &
                                (df_completo_todos_anos['data_filtro'] <= data_fim)
                            ]
                            df_ap_p = explodir_lista_json(df_p, "apreensoes")
                            if not df_ap_p.empty:
                                df_ap_p = df_ap_p[df_ap_p["INFRAÇÃO/CRIME"] != "Não se aplica"]
                            return {
                                "abordagens": int(df_p['pessoas_abordadas'].sum()) if 'pessoas_abordadas' in df_p.columns else 0,
                                "km_rodado": int(df_p['km_rodado'].sum()) if 'km_rodado' in df_p.columns else 0,
                                "apreensoes": len(df_ap_p) if not df_ap_p.empty else 0,
                                "multas": float(pd.to_numeric(df_ap_p["VALOR MULTA"], errors='coerce').fillna(0).sum()) if not df_ap_p.empty and "VALOR MULTA" in df_ap_p.columns else 0.0,
                            }

                        met_atual = metricas_do_periodo(primeiro_dia_mes_atual, hoje_cmp)
                        met_anterior = metricas_do_periodo(primeiro_dia_mes_anterior, ultimo_dia_mes_anterior)

                        cA, cB, cC, cD = st.columns(4)
                        cA.metric("Abordagens (mês atual)", met_atual["abordagens"], delta=met_atual["abordagens"] - met_anterior["abordagens"])
                        cB.metric("KM Rodado (mês atual)", f"{met_atual['km_rodado']} km", delta=met_atual["km_rodado"] - met_anterior["km_rodado"])
                        cC.metric("Apreensões (mês atual)", met_atual["apreensoes"], delta=met_atual["apreensoes"] - met_anterior["apreensoes"])
                        _delta_multas = met_atual['multas'] - met_anterior['multas']
                        _delta_multas_txt = f"{'-' if _delta_multas < 0 else '+'}R$ {abs(_delta_multas):,.2f}"
                        cD.metric("Multas (mês atual)", f"R$ {met_atual['multas']:,.2f}", delta=_delta_multas_txt)
                        st.caption(
                            f"Comparando {primeiro_dia_mes_atual.strftime('%d/%m')} até hoje, contra o mesmo intervalo do mês "
                            f"anterior ({primeiro_dia_mes_anterior.strftime('%d/%m')} a {ultimo_dia_mes_anterior.strftime('%d/%m')})."
                        )
                        
                        st.divider()
                        st.markdown("### 📋 Sequência Histórica de Relatórios Criados (Numerados)")
                        st.caption("A tabela foi dividida por item para facilitar a leitura.")

                        sub_tab01, sub_tab02, sub_tab03, sub_tab04, sub_tab05 = st.tabs([
                            "01 - Dados de Controle", "02 - Viaturas", "03 - Captura de Animais",
                            "04 - Apreensões", "05 - Estatística"
                        ])

                        with sub_tab01:
                            cols_01 = [c for c in [
                                "Nº Sequencial do Relatório", "status", "data_filtro", "unidade",
                                "finalidade", "comandante", "motorista", "data_inicial", "data_final"
                            ] if c in df.columns]
                            st.dataframe(df[cols_01], use_container_width=True)

                        with sub_tab02:
                            cols_02 = [c for c in [
                                "Nº Sequencial do Relatório", "viatura_prefixo", "km_inicial", "km_final", "km_rodado"
                            ] if c in df.columns]
                            st.dataframe(df[cols_02], use_container_width=True)

                        with sub_tab03:
                            if not df_graf_animais.empty:
                                st.dataframe(df_graf_animais, use_container_width=True)
                            else:
                                st.info("Nenhuma captura de animal registrada no período selecionado.")

                        with sub_tab04:
                            st.markdown("#### Detalhamento de Apreensões (Item 06)")
                            if not df_graf_apreensoes.empty:
                                # Traz contexto do relatório (data e comandante) para cada apreensão
                                cols_contexto = [c for c in ["Nº Sequencial do Relatório", "data_filtro", "comandante"] if c in df.columns]
                                df_contexto = df[cols_contexto].rename(columns={"Nº Sequencial do Relatório": "Nº Relatório"})
                                df_apreensoes_completo = df_graf_apreensoes.merge(df_contexto, on="Nº Relatório", how="left")
                                st.dataframe(df_apreensoes_completo, use_container_width=True)

                                col_ap_g1, col_ap_g2 = st.columns(2)
                                with col_ap_g1:
                                    st.markdown("##### Ocorrências por Tipo de Infração/Crime")
                                    df_graf1 = df_graf_apreensoes.groupby("INFRAÇÃO/CRIME")["Nº Relatório"].count().reset_index()
                                    st.bar_chart(data=df_graf1, x="INFRAÇÃO/CRIME", y="Nº Relatório")
                                with col_ap_g2:
                                    st.markdown("##### Total de Multas por Município")
                                    df_graf2 = df_graf_apreensoes.groupby("MUNICÍPIO")["VALOR MULTA"].sum().reset_index()
                                    st.bar_chart(data=df_graf2, x="MUNICÍPIO", y="VALOR MULTA")
                            else:
                                st.info("Nenhuma apreensão registrada no período selecionado.")

                        with sub_tab05:
                            if not df_graf_apreensoes.empty:
                                df_multas_rel = df_graf_apreensoes.groupby("Nº Relatório")["VALOR MULTA"].sum().reset_index()
                                df_multas_rel = df_multas_rel.rename(columns={"Nº Relatório": "Nº Sequencial do Relatório", "VALOR MULTA": "total_multas"})
                            else:
                                df_multas_rel = pd.DataFrame(columns=["Nº Sequencial do Relatório", "total_multas"])
                            df_05 = df.merge(df_multas_rel, on="Nº Sequencial do Relatório", how="left")
                            df_05["total_multas"] = df_05["total_multas"].fillna(0)
                            cols_05 = [c for c in [
                                "Nº Sequencial do Relatório", "pessoas_abordadas", "veiculos_abordados",
                                "embarcacoes_abordadas", "bo_lavrados", "autos_infracao", "total_multas"
                            ] if c in df_05.columns]
                            st.dataframe(df_05[cols_05], use_container_width=True)
                  
                        with tab_unidades:
                            st.markdown("### Comparativo Operacional: 2º Pel Miranda vs GPM Barra")

                            # Unifica os dados para trazer as unidades (Miranda, GPM Barra, etc.)
                            listagem_unidades = []
                            for idx, row in df.iterrows():
                                unid = row.get('unidade', 'Não Informada')
                                multa = pd.to_numeric(row.get('valor_multa', 0), errors='coerce')

                                # Coleta as somas terrestres deste relatório
                                soma_p_terr = 0
                                if not df_graf_terr.empty:
                                    soma_p_terr = df_graf_terr[df_graf_terr["Nº Relatório"] == row.get("Nº Sequencial do Relatório")]["Pessoas Abordadas"].sum()

                                # Coleta as somas fluviais deste relatório
                                soma_p_fluv = 0
                                if not df_graf_fluv.empty:
                                    soma_p_fluv = df_graf_fluv[df_graf_fluv["Nº Relatório"] == row.get("Nº Sequencial do Relatório")]["Pescadores Abordados"].sum()

                                listagem_unidades.append({
                                    "unidade": unid,
                                    "Total Abordados (Geral)": soma_p_terr + soma_p_fluv,
                                    "valor_multa": multa if not pd.isna(multa) else 0
                                })

                            if listagem_unidades:
                                df_unidade_nova = pd.DataFrame(listagem_unidades).groupby('unidade').sum().reset_index()
                                st.dataframe(df_unidade_nova, use_container_width=True)

                                col_un1, col_un2 = st.columns(2)
                                with col_un1:
                                    st.markdown("##### Abordagens por Unidade (Geral)")
                                    st.bar_chart(data=df_unidade_nova, x='unidade', y='Total Abordados (Geral)')
                                with col_un2:
                                    st.markdown("##### Multas Aplicadas por Unidade (R$)")
                                    st.bar_chart(data=df_unidade_nova, x='unidade', y='valor_multa')
                            else:
                                st.info("Dados de unidades não encontrados.")

                        with tab_equipes:
                            st.markdown("### Controle de Escala e Produção por Equipes (A, B e C)")

                            listagem_equipes = []
                            for idx, row in df.iterrows():
                                eqp = row.get('equipe', 'Não Informada')

                                soma_p_terr = 0
                                soma_a_terr = 0
                                if not df_graf_terr.empty:
                                    sub_t = df_graf_terr[df_graf_terr["Nº Relatório"] == row.get("Nº Sequencial do Relatório")]
                                    soma_p_terr = sub_t["Pessoas Abordadas"].sum()
                                    soma_a_terr = sub_t["Apreensões"].sum()

                                soma_p_fluv = 0
                                soma_a_fluv = 0
                                if not df_graf_fluv.empty:
                                    sub_f = df_graf_fluv[df_graf_fluv["Nº Relatório"] == row.get("Nº Sequencial do Relatório")]
                                    soma_p_fluv = sub_f["Pescadores Abordados"].sum()
                                    soma_a_fluv = sub_f["Apreensões"].sum()

                                listagem_equipes.append({
                                    "equipe": eqp,
                                    "Abordados": soma_p_terr + soma_p_fluv,
                                    "autos_infracao": soma_a_terr + soma_a_fluv
                                })

                            if listagem_equipes:
                                df_equipe_nova = pd.DataFrame(listagem_equipes).groupby('equipe').sum().reset_index()
                                st.dataframe(df_equipe_nova, use_container_width=True)
                                st.markdown("##### Eficiência de Fiscalização (Autos de Infração por Equipe)")
                                st.bar_chart(data=df_equipe_nova, x='equipe', y='autos_infracao')
                            else:
                                st.info("Dados de equipes não encontrados.")

                            st.divider()
                            st.markdown("### Produção Individual por Comandante de Guarnição")

                            listagem_comandantes = []
                            for idx, row in df.iterrows():
                                cmt = row.get('comandante', 'Não Informado')
                                nr_rel = row.get("Nº Sequencial do Relatório")

                                soma_p_terr = 0
                                if not df_graf_terr.empty:
                                    soma_p_terr = df_graf_terr[df_graf_terr["Nº Relatório"] == nr_rel]["Pessoas Abordadas"].sum()

                                soma_p_fluv = 0
                                if not df_graf_fluv.empty:
                                    soma_p_fluv = df_graf_fluv[df_graf_fluv["Nº Relatório"] == nr_rel]["Pescadores Abordados"].sum()

                                sub_ap = df_graf_apreensoes[df_graf_apreensoes["Nº Relatório"] == nr_rel] if not df_graf_apreensoes.empty else pd.DataFrame()
                                qtd_apreensoes_cmt = len(sub_ap)
                                total_multa_cmt = pd.to_numeric(sub_ap["VALOR MULTA"], errors='coerce').fillna(0).sum() if not sub_ap.empty and "VALOR MULTA" in sub_ap.columns else 0
                                total_ai_cmt = sum(len([v for v in str(x).split("; ") if v.strip()]) for x in sub_ap["AUTO INFRAÇÃO"]) if not sub_ap.empty and "AUTO INFRAÇÃO" in sub_ap.columns else 0

                                listagem_comandantes.append({
                                    "Comandante": cmt,
                                    "Abordagens (Geral)": soma_p_terr + soma_p_fluv,
                                    "Apreensões": qtd_apreensoes_cmt,
                                    "Autos de Infração": total_ai_cmt,
                                    "Total Multas (R$)": total_multa_cmt,
                                    "KM Rodado": row.get('km_rodado', 0)
                                })

                            if listagem_comandantes:
                                df_comandante_nova = pd.DataFrame(listagem_comandantes).groupby('Comandante').sum().reset_index()
                                st.dataframe(df_comandante_nova, use_container_width=True)

                                col_cmt1, col_cmt2 = st.columns(2)
                                with col_cmt1:
                                    st.markdown("##### Abordagens por Comandante")
                                    st.bar_chart(data=df_comandante_nova, x='Comandante', y='Abordagens (Geral)')
                                with col_cmt2:
                                    st.markdown("##### Multas Aplicadas por Comandante (R$)")
                                    st.bar_chart(data=df_comandante_nova, x='Comandante', y='Total Multas (R$)')

                                # --- Evolução mês a mês por comandante (todo o histórico, não só o período filtrado acima) ---
                                st.divider()
                                st.markdown("##### 📈 Evolução de Abordagens por Comandante (mês a mês)")
                                df_terr_hist = explodir_lista_json(df_completo_todos_anos, "patrulhamento_terrestre", {"PESSOAS ABORDADAS": "Pessoas Abordadas"})
                                df_fluv_hist = explodir_lista_json(df_completo_todos_anos, "patrulhamento_fluvial", {"PESCADORES ABORDADOS": "Pescadores Abordados"})

                                df_meta_hist = df_completo_todos_anos[["Nº Sequencial do Relatório", "comandante", "data_filtro"]].copy()
                                df_meta_hist["mes_ref"] = pd.to_datetime(df_meta_hist["data_filtro"]).dt.to_period("M").astype(str)

                                def _total_abordagens(nr):
                                    t = df_terr_hist[df_terr_hist["Nº Relatório"] == nr]["Pessoas Abordadas"].sum() if not df_terr_hist.empty else 0
                                    f = df_fluv_hist[df_fluv_hist["Nº Relatório"] == nr]["Pescadores Abordados"].sum() if not df_fluv_hist.empty else 0
                                    return t + f

                                df_meta_hist["abordagens"] = df_meta_hist["Nº Sequencial do Relatório"].apply(_total_abordagens)
                                df_evolucao_cmt = df_meta_hist.groupby(["mes_ref", "comandante"])["abordagens"].sum().reset_index()

                                if not df_evolucao_cmt.empty and df_evolucao_cmt["mes_ref"].nunique() > 1:
                                    df_evolucao_pivot = df_evolucao_cmt.pivot(index="mes_ref", columns="comandante", values="abordagens").fillna(0)
                                    st.line_chart(df_evolucao_pivot)
                                else:
                                    st.info("Ainda não há dados de mais de um mês para mostrar a evolução — volte aqui quando tiver relatórios de pelo menos dois meses diferentes.")
                            else:
                                st.info("Dados de comandantes não encontrados.")
            except Exception as e:
                st.error(f"Erro ao carregar o painel administrativo: {e}")
